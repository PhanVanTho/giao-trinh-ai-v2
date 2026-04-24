# -*- coding: utf-8 -*-
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- IMPORT V21.5 ENTERPRISE CLEANER ---
from dich_vu.xuat_tai_lieu.bo_loc_html import clean_for_docx as _clean_text_docx

def _set_font(run, size=13, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    # Buộc font EastAsia cũng là Times New Roman (quan trọng cho tiếng Việt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def _sanitize_xml_text(text):
    """Loại bỏ NULL bytes và control characters không hợp lệ trong XML."""
    if not text:
        return text
    # Loại bỏ NULL bytes
    text = text.replace('\x00', '')
    # Loại bỏ control characters (0x00-0x08, 0x0B, 0x0C, 0x0E-0x1F) nhưng giữ \t \n \r
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    return text

def _add_paragraph(doc, text, style=None, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    if not text:
        return None
    
    # Làm sạch markdown kỹ hơn (dùng bộ lọc V21.5)
    text = _clean_text_docx(text)
    
    # Loại bỏ NULL bytes và control characters không hợp lệ XML
    text = _sanitize_xml_text(text)
    
    text = text.strip()
    
    if not text:
        return None

    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    
    # Thụt đầu dòng 1.27cm (chuẩn giáo trình) NẾU không phải là heading hay list
    if style is None or style == "Normal":
        p.paragraph_format.first_line_indent = Inches(0.5)

    run = p.add_run(text)
    _set_font(run, size=13, bold=bold, italic=italic)
    return p

def xuat_docx(ket_qua: dict, duong_dan_docx: str):
    """
    Xuất nội dung sách ra file DOCX chuẩn định dạng giáo trình.
    """
    book = ket_qua.get("book_vi", {})
    doc = Document()

    # --- STYLE SETUP (Optional hack if styles don't exist) ---
    # (Docx mặc định đã có Normal, Heading 1...)

    # 1. Title Page
    title = book.get("title", "GIÁO TRÌNH ĐẠI HỌC")
    
    # Khoảng trắng trên
    for _ in range(5): doc.add_paragraph()
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title.upper())
    _set_font(run_title, size=24, bold=True)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tài liệu biên soạn tự động bởi hệ thống AI")
    _set_font(run_sub, size=14, italic=True)
    
    doc.add_page_break()

    # 1.5 Table of Contents (Mục Lục)
    h_toc = doc.add_heading("MỤC LỤC", level=1)
    h_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h_toc.runs:
        _set_font(run, size=16, bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    chapters_toc = book.get("chapters", [])
    for idx, ch in enumerate(chapters_toc, 1):
        # Chapter title in TOC
        toc_chap = doc.add_paragraph()
        toc_chap.paragraph_format.space_after = Pt(6)
        run_c = toc_chap.add_run(f"CHƯƠNG {idx}: {ch.get('title','').upper()}")
        _set_font(run_c, size=13, bold=True)
        
        # Sections in TOC
        for jdx, sec in enumerate(ch.get("sections", []), 1):
            toc_sec = doc.add_paragraph()
            toc_sec.paragraph_format.space_after = Pt(2)
            toc_sec.paragraph_format.left_indent = Inches(0.5) # Indent for sections
            run_s = toc_sec.add_run(f"{idx}.{jdx}. {sec.get('title','')}")
            _set_font(run_s, size=12)

    doc.add_page_break()

    # 2. Glossary (V33: Bảng thuật ngữ có định nghĩa)
    glossary = ket_qua.get("glossary", [])
    terms = ket_qua.get("terms", [])
    if glossary:
        h1 = doc.add_heading("BẢNG THUẬT NGỮ", level=1)
        for run in h1.runs:
            _set_font(run, size=16, bold=True)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for item in glossary:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run_term = p.add_run(f"{item.get('term', '')}: ")
            _set_font(run_term, size=13, bold=True)
            run_def = p.add_run(item.get('definition', ''))
            _set_font(run_def, size=13)
        doc.add_page_break()
    elif terms:
        # Fallback: danh sách thuật ngữ cũ (không có định nghĩa)
        h1 = doc.add_heading("DANH MỤC THUẬT NGỮ", level=1)
        for run in h1.runs:
            _set_font(run, size=16, bold=True)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, t in enumerate(terms[:100], 1):
            term_text = f"{i}. {t.get('term','')} - {t.get('meaning','')}"
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(term_text)
            _set_font(run, size=13)
        doc.add_page_break()

    # 3. Chapters
    chapters = book.get("chapters", [])
    for idx, ch in enumerate(chapters, 1):
        # Chapter Heading
        h1_text = f"CHƯƠNG {idx}: {ch.get('title','').upper()}"
        h1 = doc.add_heading(h1_text, level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h1.runs:
            _set_font(run, size=16, bold=True)
            # Màu đen cho chuyên nghiệp (mặc định xanh)
            run.font.color.rgb = RGBColor(0, 0, 0) 
        
        # V33: Tóm tắt chương
        summary = ch.get("summary", "")
        if summary:
            p_sum = doc.add_paragraph()
            p_sum.paragraph_format.left_indent = Inches(0.3)
            p_sum.paragraph_format.right_indent = Inches(0.3)
            p_sum.paragraph_format.space_after = Pt(12)
            run_label = p_sum.add_run("Tóm tắt chương: ")
            _set_font(run_label, size=12, bold=True, italic=True)
            run_text = p_sum.add_run(summary)
            _set_font(run_text, size=12, italic=True)

        # Sections
        for jdx, sec in enumerate(ch.get("sections", []), 1):
            # Section Heading
            h2_text = f"{idx}.{jdx}. {sec.get('title','')}"
            h2 = doc.add_heading(h2_text, level=2)
            for run in h2.runs:
                _set_font(run, size=14, bold=True)
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Content
            content = sec.get("content", "")
            if content:
                # Tách đoạn và viết
                paras = re.split(r"\n\s*\n", content)
                for raw_para in paras:
                    if raw_para.strip():
                        clean_para = raw_para.strip()
                        
                        # Xử lý Links trong DOCX (V15.0)
                        clean_para = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1 (\2)', clean_para)
                        
                        # Xử lý list item
                        if clean_para.startswith("- ") or clean_para.startswith("* "):
                            text_clean = clean_para[2:].strip()
                            _add_paragraph(doc, text_clean, style="List Bullet")
                        elif re.match(r'^\d+\.\s', clean_para):
                            # Ordered list: "1. Item"
                            text_clean = re.sub(r'^\d+\.\s', '', clean_para).strip()
                            _add_paragraph(doc, text_clean, style="List Number")
                        elif clean_para == "---":
                            # Xử lý đường kẻ ngang (V15.0)
                            p = doc.add_paragraph()
                            p_pr = p._element.get_or_add_pPr()
                            p_pbdr = qn('w:pBdr')
                            bottom = qn('w:bottom')
                            el = p_pr.find(p_pbdr)
                            if el is None:
                                el = doc.element.makeelement(p_pbdr)
                                p_pr.append(el)
                            b = doc.element.makeelement(bottom)
                            b.set(qn('w:val'), 'single')
                            b.set(qn('w:sz'), '6')
                            b.set(qn('w:space'), '1')
                            b.set(qn('w:color'), 'auto')
                            el.append(b)
                        else:
                            _add_paragraph(doc, clean_para, style="Normal")
        
        doc.add_page_break()

    # 4. Global References
    refs = ket_qua.get("references", [])
    if refs:
        h1 = doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h1.runs:
             _set_font(run, size=16, bold=True)
             run.font.color.rgb = RGBColor(0, 0, 0)
             
        # Sort by ID if possible
        try:
            refs = sorted(refs, key=lambda x: int(x.get("id", 0)) if isinstance(x, dict) else 0)
        except:
            pass

        for u in refs:
            p = doc.add_paragraph(style="Normal") # Bỏ bullet vì tự ghi [ID]
            if isinstance(u, dict):
                ref_text = f"[{u.get('id', '')}] Bách khoa toàn thư mở Wikipedia. \"{u.get('title', 'Nguồn')}\". [Trực tuyến]. Nguồn: {u.get('url', '')}. (Truy cập năm 2026)."
            else:
                ref_text = f"• {u}"
            run = p.add_run(ref_text)
            _set_font(run, size=12)

    doc.save(duong_dan_docx)
