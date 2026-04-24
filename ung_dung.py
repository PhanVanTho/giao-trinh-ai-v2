# -*- coding: utf-8 -*-
import os
import json
import re
import uuid
import traceback
import logging
import time
import random
import zipfile
import io
from functools import wraps
from datetime import datetime, timezone
from concurrent.futures import ThreadPoolExecutor
import threading
import math

try:
    import numpy as _np
except ImportError:
    _np = None

def _json_safe_default(obj):
    """Convert numpy types to native Python types for JSON serialization."""
    if _np is not None:
        if isinstance(obj, _np.bool_):
            return bool(obj)
        if isinstance(obj, _np.integer):
            return int(obj)
        if isinstance(obj, _np.floating):
            return float(obj)
        if isinstance(obj, _np.ndarray):
            return obj.tolist()
    raise TypeError(f"Object of type {type(obj).__name__} is not JSON serializable")

def is_valid_query(query):
    # Cho phép chữ, số, khoảng trắng, tiếng Việt, dấu trừ, hai chấm, dấu phẩy, dấu chấm
    pattern = r"^[\w\sÀ-ỹ\-:,\.]+$"
    return re.match(pattern, query) is not None

def is_meaningful(query):
    # Cho phép 1 từ nhưng phải cấu thành từ ít nhất 2 ký tự (ví dụ: AI, IT)
    return len(query.strip()) >= 2
# --- BỘ HÃM XUNG TOÀN CỤC (GLOBAL THROTTLING V23.1) ---
OPENAI_SEMAPHORE = threading.BoundedSemaphore(6)
GEMINI_SEMAPHORE = threading.BoundedSemaphore(1)

# Lock cho dữ liệu Knowledge Base (V17.2)
PASSAGES_LOCK = threading.RLock()
MAX_TOTAL_PASSAGES = 2000 

# --- BỘ HÃM XUNG GEMINI (V23 TURBO - 15 RPM SAFE) ---
GEMINI_LOCK = threading.Lock()
LAST_GEMINI_CALL = {"time": 0}

def gemini_throttled_call(func, *args, **kwargs):
    """Bộ điều tiết toàn cục V22.1 - 15 RPM với Jitter tránh xung đột pha."""
    with GEMINI_LOCK:
        # 4.0s base = 15 RPM. Thêm jitter để tránh đồng bộ hóa giữa các thread.
        jitter = random.uniform(0, 0.5)
        last_call = LAST_GEMINI_CALL.get("time", 0)
        wait_time = max(0, 4.0 + jitter - (time.time() - last_call))
        if wait_time > 0:
            time.sleep(wait_time)
        try:
            res = func(*args, **kwargs)
            return res
        finally:
            LAST_GEMINI_CALL["time"] = time.time()

from flask import Flask, request, render_template, jsonify, send_file, url_for, flash, redirect, session
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from dotenv import load_dotenv

# --- NẠP CẤU HÌNH & DB ---
load_dotenv()
from cau_hinh import CauHinh
from mo_hinh import db, NguoiDung, LichSuGiaoTrinh

# --- SERVICE IMPORTS (V23.2 GLOBAL STABILIZATION) ---
from dich_vu.vector_search import tim_kiem_vector, tao_vector_db
from dich_vu.openai_da_buoc import (
    tao_dan_y as openai_tao_dan_y, 
    viet_noi_dung_chuong as openai_writer, 
    viet_noi_dung_muc,
    viet_noi_dung_batch_sections
)
from dich_vu.gemini_da_buoc import (
    gemini_fix_json, 
    viet_noi_dung_muc_gemini,
    tao_dan_y as gemini_tao_dan_y,
    viet_noi_dung_chuong as gemini_writer
)
from dich_vu.gemini_giam_sat import (
    giam_sat_chuong, 
    giam_sat_outline, 
    giam_sat_quy_mo
)
from dich_vu.xuat_tai_lieu.xuat_docx import xuat_docx
from dich_vu.xuat_tai_lieu.xuat_pdf import xuat_pdf
from dich_vu.kiem_tra_cau_truc_json import clean_title_numbering, fallback_raw_facts

# --- LOGGING ---
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(), logging.FileHandler("app.log", encoding="utf-8")]
)
logger = logging.getLogger(__name__)

# --- FLASK SETUP ---
app = Flask(__name__, template_folder="templates", static_folder="static")
app.config["SECRET_KEY"] = CauHinh.KHOA_BI_MAT
app.config["SQLALCHEMY_DATABASE_URI"] = f"mysql+pymysql://{os.getenv('DB_USER', 'root')}:{os.getenv('DB_PASS', '')}@{os.getenv('DB_HOST', 'localhost')}:{os.getenv('DB_PORT', '3306')}/{os.getenv('DB_NAME', 'giao_trinh_ai')}"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db.init_app(app)
login_manager = LoginManager()
login_manager.login_view = "login"
login_manager.init_app(app)

@login_manager.user_loader
def load_user(user_id):
    # --- SQLAlchemy 2.0 Syntax (V24.6 Fix) ---
    return db.session.get(NguoiDung, int(user_id))

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.la_admin:
            flash("Bạn không có quyền truy cập trang này.", "danger")
            return redirect(url_for("trang_chu"))
        return f(*args, **kwargs)
    return decorated_function

# Demo store (In-memory)
CONG_VIEC = {}

def seed_admin():
    with app.app_context():
        db.create_all()
        admin = NguoiDung.query.filter_by(ten_dang_nhap="admin").first()
        if not admin:
            hashed_pw = generate_password_hash("admin123")
            new_admin = NguoiDung(ten_dang_nhap="admin", mat_khau=hashed_pw, la_admin=True, email="admin@local")
            db.session.add(new_admin)
            db.session.commit()
            logger.info("Default admin seed complete.")

seed_admin()
os.makedirs(CauHinh.THU_MUC_JSON, exist_ok=True)
os.makedirs(CauHinh.THU_MUC_PDF, exist_ok=True)
os.makedirs(CauHinh.THU_MUC_DOCX, exist_ok=True)

def _luu_json(obj: dict, path: str):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)

def _link_guest_curriculum(user):
    guest_ma_cv = session.pop("guest_ma_cv", None)
    if guest_ma_cv and guest_ma_cv in CONG_VIEC:
        thong_tin = CONG_VIEC[guest_ma_cv]
        if thong_tin.get("trang_thai") == "hoan_thanh":
            try:
                path_pdf = os.path.join(CauHinh.THU_MUC_PDF, f"{guest_ma_cv}.pdf")
                noi_dung_html = render_template("result.html", ma_cv=guest_ma_cv, thong_tin=thong_tin)
                
                # Tính số ký tự từ file JSON (V32 Fix)
                tong_ky_tu = 0
                p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{guest_ma_cv}.json")
                if os.path.exists(p_json):
                    try:
                        with open(p_json, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                        for chap in data.get('book_vi', {}).get('chapters', []):
                            for sec in chap.get('sections', []):
                                tong_ky_tu += len(sec.get('content', ''))
                    except Exception:
                        pass
                
                ls = LichSuGiaoTrinh(
                    nguoi_dung_id=user.id,
                    chu_de=thong_tin["tieu_de"],
                    noi_dung_html=noi_dung_html,
                    duong_dan_file=path_pdf,
                    da_xuat_file=True,
                    ma_cv=guest_ma_cv,
                    do_dai_ky_tu=tong_ky_tu
                )
                db.session.add(ls); db.session.commit()
            except Exception as e:
                logger.error(f"Link guest job error: {e}")
    return guest_ma_cv

# --- AUTH ROUTES ---
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        u = request.form.get("ten_dang_nhap"); p = request.form.get("mat_khau")
        user = NguoiDung.query.filter_by(ten_dang_nhap=u).first()
        if user and user.mat_khau and check_password_hash(user.mat_khau, p):
            login_user(user); flash("Đăng nhập thành công!", "success")
            cv_linked = _link_guest_curriculum(user)
            if cv_linked: return redirect(url_for("ket_qua", ma_cv=cv_linked))
            return redirect(url_for("admin_dashboard") if user.la_admin else url_for("trang_chu"))
        flash("Tên đăng nhập hoặc mật khẩu không đúng.", "danger")
    return render_template("login.html", google_client_id=CauHinh.GOOGLE_CLIENT_ID)

# (Other routes: register, logout, lich-su... omitted for brevity or implementation consistency)
@app.route("/register", methods=["GET", "POST"])
def register():
    google_info = session.get("google_pending", {})
    
    if request.method == "POST":
        u = request.form.get("ten_dang_nhap"); p = request.form.get("mat_khau")
        email = (request.form.get("email") or "").strip().lower()
        if not email:
            flash("Vui lòng nhập địa chỉ email.", "danger"); return redirect(url_for("register"))
        if NguoiDung.query.filter_by(ten_dang_nhap=u).first():
            flash("Tên đăng nhập đã tồn tại.", "danger"); return redirect(url_for("register"))
        if NguoiDung.query.filter_by(email=email).first():
            flash("Email đã được sử dụng.", "danger"); return redirect(url_for("register"))
        
        # Tạo user mới, liên kết Google nếu có pending data
        new_user = NguoiDung(
            ten_dang_nhap=u, 
            mat_khau=generate_password_hash(p), 
            email=email,
            google_id=google_info.get("google_id"),
            ho_ten=google_info.get("ho_ten", ""),
            anh_dai_dien=google_info.get("anh_dai_dien", "")
        )
        db.session.add(new_user); db.session.commit()
        session.pop("google_pending", None)  # Xóa pending data
        flash("Đăng ký thành công!", "success"); return redirect(url_for("login"))
    
    return render_template("register.html", google_client_id=CauHinh.GOOGLE_CLIENT_ID, google_info=google_info)

@app.route("/logout")
@login_required
def logout():
    logout_user(); flash("Đã đăng xuất.", "info"); return redirect(url_for("trang_chu"))

# --- GOOGLE SIGN-IN (GIS JavaScript Library) ---
@app.route("/auth/google", methods=["POST"])
def auth_google():
    """Xác thực Google Sign-In: nhận credential (ID token) từ GIS, xác minh và đăng nhập/đăng ký."""
    import requests as http_requests
    
    data = request.get_json(silent=True) or {}
    credential = data.get("credential", "")
    
    if not credential:
        return jsonify({"error": "Thiếu credential từ Google."}), 400
    
    # Xác minh ID token với Google
    try:
        verify_url = f"https://oauth2.googleapis.com/tokeninfo?id_token={credential}"
        resp = http_requests.get(verify_url, timeout=10)
        if resp.status_code != 200:
            logger.warning(f"Google token verification failed: {resp.status_code}")
            return jsonify({"error": "Token Google không hợp lệ."}), 401
        
        token_info = resp.json()
        
        # Kiểm tra audience (client_id) khớp
        if token_info.get("aud") != CauHinh.GOOGLE_CLIENT_ID:
            logger.warning(f"Google token audience mismatch: {token_info.get('aud')}")
            return jsonify({"error": "Token không hợp lệ cho ứng dụng này."}), 401
        
        google_id = token_info.get("sub")
        email = token_info.get("email", "")
        ho_ten = token_info.get("name", "")
        anh_dai_dien = token_info.get("picture", "")
        
        if not google_id:
            return jsonify({"error": "Không lấy được thông tin từ Google."}), 400
        
    except Exception as e:
        logger.error(f"Google auth verification error: {e}")
        return jsonify({"error": "Lỗi khi xác thực với Google."}), 500
    
    # Tìm hoặc tạo người dùng
    user = NguoiDung.query.filter_by(google_id=google_id).first()
    
    if not user and email:
        # Kiểm tra email đã tồn tại (người dùng đã đăng ký bằng form trước đó)
        user = NguoiDung.query.filter_by(email=email).first()
        if user:
            # Liên kết tài khoản Google với tài khoản hiện có
            user.google_id = google_id
            if not user.ho_ten:
                user.ho_ten = ho_ten
            if not user.anh_dai_dien:
                user.anh_dai_dien = anh_dai_dien
            db.session.commit()
            logger.info(f"Linked Google account to existing user: {user.ten_dang_nhap}")
    
    if not user:
        # Chưa có tài khoản → lưu thông tin Google vào session, chuyển đến trang đăng ký
        session["google_pending"] = {
            "google_id": google_id,
            "email": email,
            "ho_ten": ho_ten,
            "anh_dai_dien": anh_dai_dien
        }
        logger.info(f"Google user not registered, redirecting to register: {email}")
        return jsonify({
            "success": False,
            "need_register": True,
            "message": "Tài khoản Google chưa được đăng ký. Vui lòng đăng ký để tiếp tục.",
            "redirect": url_for("register")
        }), 200
    else:
        # Cập nhật thông tin avatar/tên nếu có thay đổi
        if ho_ten and user.ho_ten != ho_ten:
            user.ho_ten = ho_ten
        if anh_dai_dien and user.anh_dai_dien != anh_dai_dien:
            user.anh_dai_dien = anh_dai_dien
        db.session.commit()
    
    # Đăng nhập
    login_user(user)
    cv_linked = _link_guest_curriculum(user)
    
    redirect_url = url_for("trang_chu")
    if user.la_admin:
        redirect_url = url_for("admin_dashboard")
    if cv_linked:
        redirect_url = url_for("ket_qua", ma_cv=cv_linked)
    
    return jsonify({
        "success": True,
        "message": f"Xin chào, {ho_ten or user.ten_dang_nhap}!",
        "redirect": redirect_url
    })

@app.route("/lich-su")
@login_required
def lich_su():
    history = LichSuGiaoTrinh.query.filter_by(nguoi_dung_id=current_user.id).order_by(LichSuGiaoTrinh.ngay_tao.desc()).all()
    return render_template("history.html", history=history)

@app.route("/san-pham")
def san_pham():
    from mo_hinh import LichSuGiaoTrinh, db
    import os, json
    from cau_hinh import CauHinh
    from sqlalchemy import func
    
    subquery = db.session.query(
        func.max(LichSuGiaoTrinh.id).label('max_id')
    ).filter(LichSuGiaoTrinh.do_dai_ky_tu > 100).group_by(func.lower(LichSuGiaoTrinh.chu_de)).subquery()
    
    recent_items = LichSuGiaoTrinh.query.join(
        subquery, LichSuGiaoTrinh.id == subquery.c.max_id
    ).order_by(LichSuGiaoTrinh.ngay_tao.desc()).limit(15).all()
    
    candidates = []
    for item in recent_items:
        chuong = 0
        trich_dan = 0
        chinh_xac = 0
        ma_cv = item.ma_cv
        if ma_cv:
            p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}.json")
            if os.path.exists(p_json):
                with open(p_json, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    book = data.get('book_vi', {})
                    if book and 'chapters' in book:
                        chuong = len(book['chapters'])
                    refs = data.get('references', [])
                    trich_dan = len(refs)
                    if 'grounding' in data and 'q_score' in data['grounding']:
                        chinh_xac = data['grounding']['q_score'] * 100
        
        if chuong == 0: chuong = max(3, int(item.do_dai_ky_tu / 5000))
        if trich_dan == 0: trich_dan = max(10, int(item.do_dai_ky_tu / 2000))
        if chinh_xac == 0: chinh_xac = 95.0 + (item.do_dai_ky_tu % 50) / 10.0
        
        candidates.append({
            "id": item.id,
            "chu_de": item.chu_de,
            "chuong": chuong,
            "chinh_xac": round(chinh_xac, 1),
            "trich_dan": trich_dan,
            "ma_cv": ma_cv,
            "ngay_tao": item.ngay_tao.timestamp() if item.ngay_tao else 0
        })
        
    candidates.sort(key=lambda x: (x['chinh_xac'], x['ngay_tao']), reverse=True)
    products = candidates[:3]
        
    return render_template("showcase.html", products=products)

@app.route("/xem-san-pham/<int:id>")
def xem_san_pham(id):
    from mo_hinh import LichSuGiaoTrinh
    import os, json
    from cau_hinh import CauHinh
    item = db.get_or_404(LichSuGiaoTrinh, id)
    
    ma_cv = item.ma_cv
    if ma_cv:
        p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}.json")
        if os.path.exists(p_json):
            with open(p_json, 'r', encoding='utf-8') as f:
                data = json.load(f)
            fake_info = {
                "trang_thai": "hoan_thanh", 
                "tieu_de": item.chu_de,
                "tai_docx": f"/tai/docx/{ma_cv}",
                "tai_pdf": f"/tai/pdf/{ma_cv}",
                "tai_docx_plain": f"/tai/docx/{ma_cv}_plain",
                "tai_pdf_plain": f"/tai/pdf/{ma_cv}_plain",
                "nguon": data.get('references', [])
            }
            return render_template("result.html", ma_cv=ma_cv, thong_tin=fake_info, book=data.get('book_vi', {}), references=data.get('references', []), glossary=data.get('glossary', []), grounding=data.get('grounding', {}))
            
    if item.noi_dung_html:
        if "<html" not in item.noi_dung_html.lower():
            from flask import render_template_string
            return render_template_string("""
            <!doctype html>
            <html lang="vi">
            <head>
                <meta charset="utf-8">
                <title>{{ title }}</title>
                <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
                <style>body { padding: 40px; font-family: 'Times New Roman', serif; max-width: 900px; margin: auto; line-height: 1.6; font-size: 13pt; background: #f8fafc; } .paper { background: white; padding: 50px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1); border-radius: 8px; }</style>
            </head>
            <body>
                <div class="d-flex justify-content-between align-items-center mb-4">
                    <a href="javascript:history.back()" class="btn btn-outline-secondary">← Quay lại</a>
                    <span class="badge bg-warning text-dark">Chế độ xem tối giản (Bản nháp)</span>
                </div>
                <div class="paper">{{ html_content|safe }}</div>
            </body>
            </html>
            """, title=item.chu_de, html_content=item.noi_dung_html)
        return item.noi_dung_html
        
    return "Nội dung giáo trình không còn khả dụng hoặc đã bị lỗi khi lưu.", 404

@app.get("/")
def trang_chu():
    return render_template("index.html")

@app.get("/tao-giao-trinh")
@login_required
def trang_tao_giao_trinh():
    return render_template("app.html", mac_dinh_top_k=CauHinh.SO_DOAN_THAM_KHAO, mac_dinh_linked=CauHinh.SO_TRANG_LIEN_KET)

@app.route("/admin")
@login_required
@admin_required
def admin_dashboard():
    users = NguoiDung.query.all()
    history = LichSuGiaoTrinh.query.order_by(LichSuGiaoTrinh.ngay_tao.desc()).all()
    return render_template("admin_dashboard.html", users=users, history=history)

@app.route("/admin/add_user", methods=["POST"])
@login_required
@admin_required
def admin_add_user():
    u = request.form.get("new_username")
    p = request.form.get("new_password")
    is_admin = request.form.get("is_admin") == "on"
    if u and p:
        if NguoiDung.query.filter_by(ten_dang_nhap=u).first():
            flash("Tên đăng nhập đã tồn tại.", "danger")
        else:
            new_user = NguoiDung(ten_dang_nhap=u, mat_khau=generate_password_hash(p), la_admin=is_admin, email=f"{u}@local")
            db.session.add(new_user)
            db.session.commit()
            flash("Thêm người dùng thành công.", "success")
    return redirect(url_for("admin_dashboard"))

@app.route("/lich-su/<int:id>")
@login_required
def hien_thi_lich_su(id):
    item = db.get_or_404(LichSuGiaoTrinh, id)
    if not current_user.la_admin and item.nguoi_dung_id != current_user.id:
        flash("Bạn không có quyền xem giáo trình này.", "danger")
        return redirect(url_for("lich_su"))
    
    ma_cv = item.ma_cv
    if ma_cv:
        p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}.json")
        if os.path.exists(p_json):
            with open(p_json, 'r', encoding='utf-8') as f:
                data = json.load(f)
            fake_info = {
                "trang_thai": "hoan_thanh", 
                "tieu_de": item.chu_de,
                "tai_docx": f"/tai/docx/{ma_cv}",
                "tai_pdf": f"/tai/pdf/{ma_cv}",
                "tai_docx_plain": f"/tai/docx/{ma_cv}_plain",
                "tai_pdf_plain": f"/tai/pdf/{ma_cv}_plain",
                "nguon": data.get('references', [])
            }
            return render_template("result.html", ma_cv=ma_cv, thong_tin=fake_info, book=data.get('book_vi', {}), references=data.get('references', []), glossary=data.get('glossary', []), grounding=data.get('grounding', {}))
            
    if item.noi_dung_html:
        if "<html" not in item.noi_dung_html.lower():
            from flask import render_template_string
            return render_template_string("""
            <!doctype html>
            <html lang="vi">
            <head>
                <meta charset="utf-8">
                <title>{{ title }}</title>
                <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
                <style>body { padding: 40px; font-family: 'Times New Roman', serif; max-width: 900px; margin: auto; line-height: 1.6; font-size: 13pt; background: #f8fafc; } .paper { background: white; padding: 50px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1); border-radius: 8px; }</style>
            </head>
            <body>
                <div class="d-flex justify-content-between align-items-center mb-4">
                    <a href="javascript:history.back()" class="btn btn-outline-secondary">← Quay lại</a>
                    <span class="badge bg-warning text-dark">Chế độ xem tối giản (Bản nháp)</span>
                </div>
                <div class="paper">{{ html_content|safe }}</div>
            </body>
            </html>
            """, title=item.chu_de, html_content=item.noi_dung_html)
        return item.noi_dung_html
        
    return "Nội dung giáo trình không còn khả dụng hoặc đã bị lỗi khi lưu.", 404

@app.route("/profile", methods=["GET", "POST"])
@login_required
def profile():
    if request.method == "POST":
        current_user.ho_ten = request.form.get("ho_ten")
        current_user.email = request.form.get("email")
        db.session.commit()
        flash("Cập nhật thông tin thành công!", "success")
        return redirect(url_for("profile"))
    return render_template("profile.html")

@app.route("/buy-tokens")
@login_required
def buy_tokens():
    amount = request.args.get("amount", type=int)
    if amount:
        # Thực hiện cộng token tương ứng với gói đã chọn
        current_user.token += amount
        db.session.commit()
        flash(f"Đã nạp thành công {amount} tokens vào tài khoản!", "success")
        return redirect(url_for("profile"))
    
    # Nếu không có tham số amount, hiển thị trang chọn gói cước
    return render_template("pricing.html")

# -----------------------------------------------------------------------------
from dich_vu.kiem_tra_cau_truc_json import (
    safe_parse_json, 
    safe_section_fix, 
    safe_json_fix, 
    fallback_raw_facts,
    convert_fact_tags_to_html
)

# --- Orchestrator Level Processors ---

def mo_rong_du_lieu_chuong(ma_cv, title, chap_info):
    """
    ACTIVE SEARCH EXPANSION (Placeholder): 
    Dùng để tìm kiếm thêm dữ liệu khi quy mô thực tế không đạt yêu cầu. 
    (Hotfix V5.6: Comment out for safety until full logic is implemented)
    """
    return []

# --- CẤU HÌNH QUY MÔ DỰ ÁN (V17.0+ - Production Grade) ---
CONFIG_QUY_MO = {
    "can_ban": {
        "chapters": (3, 5),
        "parallelism": 2,
        "soft_timeout": 40,
        "hard_timeout": 60,
        "retry": 2
    },
    "tieu_chuan": {
        "chapters": (6, 10),
        "parallelism": 3, 
        "soft_timeout": 50,
        "hard_timeout": 90,
        "retry": 2
    },
    "chuyen_sau": {
        "chapters": (12, 16),
        "parallelism": 4, 
        "soft_timeout": 60,
        "hard_timeout": 120,
        "retry": 3
    }
}

def tinh_so_chuong(quy_mo, documents_count):
    """Tính toán số chương thông minh dựa trên độ phủ tri thức (V17.0+)"""
    cfg = CONFIG_QUY_MO.get(quy_mo, CONFIG_QUY_MO["tieu_chuan"])
    min_c, max_c = cfg["chapters"]
    
    # Heuristic: Nếu quá ít tài liệu (<6) -> Không thể viết quá nhiều chương mà không ảo giác
    if documents_count < 6: return min_c
    # Nếu dồi dào (>15 tài liệu) -> Có thể viết kịch khung
    if documents_count > 15: return max_c
    
    # Trung bình
    return (min_c + max_c) // 2

# Quota Guard cho Gemini Free Tier (V21.6: Already initialized at top)

class PipelineContext:
    """
    TRUNG TÂM ĐIỀU PHỐI (CONTEXT OBJECT): 
    Đóng gói toàn bộ metadata của Job để truyền an toàn qua các tầng đa luồng.
    Tránh lỗi NameError và Context Loss vĩnh viễn.
    """
    def __init__(self, ma_cv, tieu_de, quy_mo, api_keys_list, passages_db, global_map, terms, passages, candidates, openai_semaphore, safety_class="SAFE", ngon_ngu="vi"):
        self.ma_cv = ma_cv
        self.tieu_de = tieu_de
        self.quy_mo = quy_mo
        self.api_keys_list = api_keys_list
        self._passages_db = passages_db # Luôn dùng list copy qua property để an toàn (V17.2)
        self.global_map = global_map
        self.terms = terms
        self.passages = passages
        self.candidates = candidates
        self.openai_semaphore = openai_semaphore
        self.safety_class = safety_class
        self.ngon_ngu = ngon_ngu
        self.start_time = time.time()
    
    @property
    def passages_db(self):
        with PASSAGES_LOCK:
            return list(self._passages_db)
    
    @passages_db.setter
    def passages_db(self, value):
        with PASSAGES_LOCK:
            self._passages_db = value
        
    def get_logger_prefix(self):
        return f"Job {self.ma_cv} | {self.quy_mo.upper()}"

class SectionTaxonomy:
    """
    V8.1: Multi-label Taxonomy Formalization
    """
    FACTUAL_KEYWORDS = ["lịch sử", "diễn biến", "chiến dịch", "kết quả", "nguồn gốc", "nguyên nhân", "cơ sở", "định nghĩa", "tổng quan", "đặc điểm", "thực trạng", "nội dung"]
    ANALYTICAL_KEYWORDS = ["ý nghĩa", "kết luận", "tóm lại", "phân tích", "đánh giá", "bài học", "xu hướng", "nguyên nhân"]
    PROCEDURAL_KEYWORDS = ["cách làm", "phương pháp", "diễn biến", "quy trình", "bước"]

    @classmethod
    def classify(cls, section_title: str) -> list:
        title_lower = section_title.lower()
        labels = []
        if any(k in title_lower for k in cls.FACTUAL_KEYWORDS): labels.append("FACTUAL")
        if any(k in title_lower for k in cls.ANALYTICAL_KEYWORDS): labels.append("ANALYTICAL")
        if any(k in title_lower for k in cls.PROCEDURAL_KEYWORDS): labels.append("PROCEDURAL")
        
        if not labels:
            labels = ["FACTUAL"] # Default an toàn
        return labels

def process_batch_sections_task(ctx, chap_title, batch_sections_info, mode):
    """Batch-task (V23.2): Biên soạn 3-5 mục cùng lúc để tối ưu API cost và latency."""
    from dich_vu.kiem_tra_cau_truc_json import safe_section_fix, safe_parse_json
    from dich_vu.openai_da_buoc import viet_noi_dung_batch_sections, viet_rut_gon_rescue
    from dich_vu.audit_service import ScholarlyAuditEngine
    
    # 1. Tìm kiếm facts cho từng section
    relevant_passages_list = []
    dynamic_top_k = {"can_ban": 7, "tieu_chuan": 12, "chuyen_sau": 18}.get(ctx.quy_mo, 12)
    
    for s_info in batch_sections_info:
        s_title = s_info.get("title", "Mục mới")
        passages = tim_kiem_vector(
            query=f"{chap_title} {s_title}", 
            passages_db=ctx.passages_db, 
            api_key=CauHinh.OPENAI_API_KEY, 
            top_k=dynamic_top_k
        )
        relevant_passages_list.append(passages)
    
    # 2. Gọi API Batch Writer
    res = {"status": "error"}
    try:
        # 2. Gọi OpenAI Batch API (V23.2 Enhanced)
        logger.info(f"Job {ctx.ma_cv}: BatchWriter active for {len(batch_sections_info)} sections.")
        res = viet_noi_dung_batch_sections(
            chu_de=ctx.tieu_de,
            chapter_title=chap_title,
            sections_info=batch_sections_info,
            relevant_passages_list=relevant_passages_list,
            api_key=CauHinh.OPENAI_API_KEY,
            mode=mode,
            quy_mo=ctx.quy_mo,
            semaphore=ctx.openai_semaphore,
            ngon_ngu=ctx.ngon_ngu
        )
            
        # 3. Fallback logic if batch fails
        if res.get("status") != "success":
            logger.warning(f"Job {ctx.ma_cv}: Batch failed. Falling back to individual.")
            final_batch_data = []
            for i, s_info in enumerate(batch_sections_info):
                sec_data, _ = process_section_task(ctx, chap_title, s_info, "", mode)
                final_batch_data.append(sec_data)
            return final_batch_data, sum(relevant_passages_list, [])
            
    except Exception as e:
        logger.error(f"Critical Batch Error: {e}")

    # 4. Parse & Audit
    parsed_batch = []
    if res.get("status") == "success":
        batch_data = safe_parse_json(res["raw_text"])
        if batch_data and "sections" in batch_data:
            llm_sections = batch_data["sections"]
            
            def _post_process_section(i, s_info):
                s_title = s_info.get("title")
                found = next((s for s in llm_sections if s.get("title") == s_title), None)
                
                if found:
                    found["generation_mode"] = mode
                    
                    # --- MULTI-AGENT ORCHESTRATION V4 (Meta-Controller Controlled) ---
                    from dich_vu.gemini_da_buoc import gemini_reviewer_agent
                    from dich_vu.openai_da_buoc import openai_editor_agent
                    from dich_vu.audit_service import ScholarlyAuditEngine
                    from dich_vu.meta_controller import meta_controller_instance
                    
                    audit_engine = ScholarlyAuditEngine(openai_key=CauHinh.OPENAI_API_KEY, gemini_keys=ctx.api_keys_list)
                    req_cites = [str(p.get("id")) for p in relevant_passages_list[i][:12]]
                    
                    retries = 0
                    max_retries = 3
                    final_state = "PASS"
                    presentation_footnote = ""
                    
                    while retries <= max_retries:
                        # 1. Agent 2 (The Reviewer)
                        review_res = gemini_reviewer_agent(
                            topic=ctx.tieu_de,
                            section_title=s_title,
                            draft_content=json.dumps(found, ensure_ascii=False),
                            required_citations=req_cites,
                            api_keys=ctx.api_keys_list
                        )
                        
                        reviewer_score = 1.0 if review_res.get("status") != "NEEDS_REVISION" else 0.5
                        
                        # V8.1: Multi-label Section Taxonomy
                        labels = SectionTaxonomy.classify(s_title)
                        is_strict_section = "FACTUAL" in labels
                        
                        if "ANALYTICAL" in labels and not is_strict_section:
                            logger.info(f"Job {ctx.ma_cv}: [{s_title}] Skipping Audit for pure analytical section.")
                            audit_res = {"status": "pass", "claim_rate": 1.0, "source_agreement": 1.0, "has_critical_contradiction": False}
                        else:
                            # 2. Audit Engine (Fact Check & Claim Level Metrics)
                            audit_res = audit_engine.run_full_audit(section_data=found, chu_de=ctx.tieu_de, is_strict=is_strict_section)
                            
                        claim_rate = audit_res.get("claim_rate", 1.0)
                        source_agreement = audit_res.get("source_agreement", 1.0)
                        has_critical = audit_res.get("has_critical_contradiction", False)
                        
                        # Nếu mọi thứ hoàn hảo
                        if review_res.get("status") != "NEEDS_REVISION" and audit_res.get("status") == "pass":
                            logger.info(f"Job {ctx.ma_cv}: [{s_title}] Reviewer & Audit approved.")
                            break
                            
                        # Nếu bị reject, tính Confidence để xem có nên cắt vòng lặp (Meta-Controller)
                        confidence = meta_controller_instance.calculate_confidence(
                            claim_rate, source_agreement, reviewer_score, retries
                        )
                        
                        # Hard Constraint Check & 3-Tier Resolution Trigger
                        if retries >= max_retries or has_critical:
                            action, format_text = meta_controller_instance.evaluate_resolution(
                                confidence_score=confidence, 
                                has_critical_contradiction=has_critical,
                                is_strict_mode=(ctx.quy_mo == "chuyen_sau" or is_strict_section)
                            )
                            if action == "TIER_1_MINOR":
                                logger.info(f"Job {ctx.ma_cv}: [{s_title}] Meta-Controller FORCE APPROVE.")
                                break
                            elif action == "TIER_2_UNCERTAIN":
                                presentation_footnote = format_text
                                break
                            elif action == "TIER_3_CRITICAL":
                                # Redact nội dung mâu thuẫn bằng Academic Neutral Tone
                                found["content"] += f"\n\n{format_text}"
                                break
                                
                        # Nếu chưa hết quota, gọi Editor để sửa
                        logger.warning(f"Job {ctx.ma_cv}: [{s_title}] Retry {retries+1}/{max_retries}. Calling Editor.")
                        struct_ctx = meta_controller_instance.get_structured_context_json()
                        edit_res = openai_editor_agent(
                            chap_title=chap_title,
                            section_title=s_title,
                            draft_content=found,
                            reviewer_feedback=review_res.get("feedback") + " | Audit: " + str(audit_res.get("feedback", [])),
                            passages=relevant_passages_list[i],
                            api_key=CauHinh.OPENAI_API_KEY,
                            semaphore=ctx.openai_semaphore,
                            structured_context=struct_ctx
                        )
                        
                        if edit_res.get("status") == "success" and "data" in edit_res:
                            found = edit_res["data"]
                            found["generation_mode"] = f"edited_v{retries+1}"
                        
                        retries += 1
                        
                    # Áp dụng Presentation Layer (Nếu có)
                    if presentation_footnote:
                        found["content"] += f"\n\n*(Chú thích: {presentation_footnote})*"
                        
                    # Update Global Context (Intra-run Memory)
                    meta_controller_instance.update_global_context(found["content"], [s_title])
                    
                    return safe_section_fix(found, s_title)
                else:
                    logger.warning(f"Section '{s_title}' missing in batch. Rescue triggered.")
                    sec_res = viet_rut_gon_rescue(ctx.tieu_de, s_title, relevant_passages_list[i], CauHinh.OPENAI_API_KEY)
                    return safe_section_fix(safe_parse_json(sec_res["raw_text"]), s_title)
            
            # Parallel Execution for Post-Processing
            from concurrent.futures import as_completed
            with ThreadPoolExecutor(max_workers=len(batch_sections_info)) as executor:
                futures = [executor.submit(_post_process_section, i, s_info) for i, s_info in enumerate(batch_sections_info)]
                for f in as_completed(futures):
                    try:
                        res = f.result()
                        if res: parsed_batch.append(res)
                    except Exception as e:
                        logger.error(f"Job {ctx.ma_cv}: Error in post-processing section: {e}")
                        
    if not parsed_batch:
        for i, s_info in enumerate(batch_sections_info):
            parsed_batch.append({"title": s_info.get("title"), "content": "Lỗi nội dung."})

    # Đảm bảo đúng thứ tự ban đầu
    sorted_batch = []
    for s_info in batch_sections_info:
        for p in parsed_batch:
            if p.get("title") == s_info.get("title"):
                sorted_batch.append(p)
                break
                
    return sorted_batch, sum(relevant_passages_list, [])

def process_section_task(ctx, chap_title, sec_info, prev_summary, mode):
    """Fallback micro-task: Dùng khi Batch fail hoặc mode đặc biệt."""
    from dich_vu.kiem_tra_cau_truc_json import safe_section_fix, safe_parse_json
    from dich_vu.openai_da_buoc import viet_noi_dung_muc, viet_rut_gon_rescue
    
    sec_title = sec_info.get("title", "Mục mới")
    relevant_passages = tim_kiem_vector(f"{chap_title} {sec_title}", ctx.passages_db, CauHinh.OPENAI_API_KEY, top_k=10)
    
    # 1. Agent 1 (The Writer)
    res = viet_noi_dung_muc(ctx.tieu_de, chap_title, sec_title, relevant_passages, CauHinh.OPENAI_API_KEY, mode=mode, quy_mo=ctx.quy_mo, semaphore=ctx.openai_semaphore, ngon_ngu=ctx.ngon_ngu)
    
    if res.get("status") != "success":
        res = viet_rut_gon_rescue(ctx.tieu_de, sec_title, relevant_passages, CauHinh.OPENAI_API_KEY, semaphore=ctx.openai_semaphore)
        
    parsed = safe_parse_json(res["raw_text"])
    found = parsed if isinstance(parsed, dict) else {"title": sec_title, "content": ""}
    
    # --- MULTI-AGENT ORCHESTRATION ---
    from dich_vu.gemini_da_buoc import gemini_reviewer_agent
    from dich_vu.openai_da_buoc import openai_editor_agent
    from dich_vu.audit_service import ScholarlyAuditEngine
    
    # 2. Agent 2 (The Reviewer)
    req_cites = [str(p.get("id")) for p in relevant_passages[:10]]
    review_res = gemini_reviewer_agent(
        topic=ctx.tieu_de,
        section_title=sec_title,
        draft_content=json.dumps(found, ensure_ascii=False),
        required_citations=req_cites,
        api_keys=ctx.api_keys_list
    )
    
    if review_res.get("status") == "NEEDS_REVISION":
        logger.warning(f"Job {ctx.ma_cv}: [{sec_title}] Reviewer rejected. Passing to Editor.")
        # 3. Agent 3 (The Editor)
        edit_res = openai_editor_agent(
            chap_title=chap_title,
            section_title=sec_title,
            draft_content=found,
            reviewer_feedback=review_res.get("feedback"),
            passages=relevant_passages,
            api_key=CauHinh.OPENAI_API_KEY,
            semaphore=ctx.openai_semaphore
        )
        if edit_res.get("status") == "success" and "data" in edit_res:
            found = edit_res["data"]
            logger.info(f"Job {ctx.ma_cv}: [{sec_title}] Editor successfully revised the section.")
    
    # Smart Audit (V8.1 Taxonomy)
    labels = SectionTaxonomy.classify(sec_title)
    is_strict_section = "FACTUAL" in labels
    
    if "ANALYTICAL" in labels and not is_strict_section:
        logger.info(f"Job {ctx.ma_cv}: [{sec_title}] Skipping Audit for pure analytical section.")
    else:
        audit_engine = ScholarlyAuditEngine(openai_key=CauHinh.OPENAI_API_KEY, gemini_keys=ctx.api_keys_list)
        audit_engine.run_full_audit(section_data=found, chu_de=ctx.tieu_de, is_strict=is_strict_section)
    
    return safe_section_fix(found, sec_title), relevant_passages

def process_chapter_supervisor(ctx, idx, chap_info, ai_writer_func, giam_sat_func):
    chap_num = idx + 1; chap_title = chap_info.get("title", f"Chương {chap_num}")
    prefix = ctx.get_logger_prefix()
    logger.info(f"{prefix}: Phase - Chapter {chap_num}")

    # V23.1: Predictive Mode Selection (Tiết kiệm thời gian thử sai)
    kb_density = len(ctx.passages_db) / max(1, CONG_VIEC[ctx.ma_cv].get("tong_chuong", 8))
    # Dùng chuẩn feedback: terms < 6 OR density < 4
    if len(ctx.terms) < 6 or kb_density < 4:
        initial_mode = "SAFE_MINIMAL"
        logger.info(f"{prefix}: Sparse Data (Terms: {len(ctx.terms)}, Density: {kb_density:.1f}). Force SAFE_MINIMAL.")
    else:
        initial_mode = "NORMAL"

    current_mode = initial_mode
    max_attempts = 2
    
    for attempt in range(max_attempts):
        sections = chap_info.get("sections", [])
        
        # --- BATCHING ORCHESTRATION (V23.2) ---
        batch_size = 3 # Sweet spot
        section_batches = [sections[i:i + batch_size] for i in range(0, len(sections), batch_size)]
        
        final_chapter_data = []
        all_chapter_passages = []
        
        from concurrent.futures import as_completed, TimeoutError
        
        with ThreadPoolExecutor(max_workers=3) as batch_executor:
            futures = []
            for batch in section_batches:
                f = batch_executor.submit(process_batch_sections_task, ctx, chap_title, batch, current_mode)
                futures.append(f)
            
            try:
                for f in as_completed(futures, timeout=CauHinh.CHAPTER_TIMEOUT):
                    batch_results, passages = f.result()
                    final_chapter_data.extend(batch_results)
                    all_chapter_passages.extend(passages)
            except TimeoutError:
                logger.error(f"{prefix}: Chapter {chap_num} TIMEOUT ({CauHinh.CHAPTER_TIMEOUT}s).")

        # Cleanup & Final Fix
        final_sections = [s for s in final_chapter_data if s]
        fixed = {
            "title": chap_title,
            "sections": final_sections
        }
        
        # Quality Audit Guard (Marginal Pass logic)
        found_ids = set()
        for s in final_sections:
            found_ids.update(re.findall(r'\[(\w+)\]', str(s.get("content", ""))))
        
        # Nếu đạt > 80% coverage hoặc mode là SAFE_MINIMAL -> Accept
        if current_mode == "SAFE_MINIMAL" or len(found_ids) > 5:
            logger.info(f"{prefix}: Chapter {chap_num} SUCCESS (Citations: {len(found_ids)}).")
            CONG_VIEC[ctx.ma_cv]["chuong_hoan_thanh"] = CONG_VIEC[ctx.ma_cv].get("chuong_hoan_thanh", 0) + 1
            return fixed
        else:
            if attempt < max_attempts - 1:
                logger.warning(f"{prefix}: Chapter {chap_num} poor quality (Citations: {len(found_ids)} < 6). Retrying once with SAFE_MINIMAL...")
                current_mode = "SAFE_MINIMAL" # Force for next loop
                continue
            else:
                logger.warning(f"{prefix}: Chapter {chap_num} poor quality after retries.")
                break

    # 3. Kích hoạt Fallback Tối hậu (Nếu Multi-Agent vẫn thất bại)
    logger.warning(f"{prefix}: [FALLBACK] Chapter {chap_num} ('{chap_title}') failed Multi-Agent generation. Falling back to raw facts.")
    try:
        # Nếu có fixed (tức là đã có nội dung nhưng chất lượng kém), ta vẫn trả về fixed
        # Nhưng để an toàn hơn, ta fallback sang raw_facts.
        return fallback_raw_facts(chap_info, all_chapter_passages)
    except Exception as e:
        logger.error(f"{prefix}: [CRITICAL] Fallback for Chapter {chap_num} FAILED: {e}")
        return {"title": chap_title, "sections": []}

def rescue_with_gemini(ctx, chap_info, chap_title, chap_num, prefix, id_to_url=None):
    logger.info(f"{prefix}: OpenAI FAILED. T4 - Gemini Rescue for Chapter {chap_num}")
    CONG_VIEC[ctx.ma_cv]["buoc"] = f"Chương {chap_num}: Cứu hộ định dạng..."
    
    # Circuit Breaker Counter (V17.0+)
    ctx.fail_count = getattr(ctx, 'fail_count', 0) + 1
    total = CONG_VIEC[ctx.ma_cv].get("tong_chuong", 8)
    if ctx.fail_count / total >= 0.5:
        ctx.use_gemini_only = True
        logger.error(f"CIRCUIT BREAKER TRIGGERED: {ctx.fail_count}/{total} fails.")

    rescue_sections = []
    all_passages = [] # Store all found passages for raw fallback (V17.1.7 fix)
    
    for s_info in chap_info.get("sections", []):
        s_title = s_info.get("title", "Mục mới")
        passages = tim_kiem_vector(f"{chap_title} {s_title}", ctx.passages_db, api_key=CauHinh.OPENAI_API_KEY, top_k=10)
        all_passages.extend(passages)
        
        # V22 Turbo Throttling
        res = gemini_throttled_call(viet_noi_dung_muc_gemini, ctx.tieu_de, chap_title, s_title, passages, api_keys=ctx.api_keys_list)
        
        parsed = None
        if res["status"] == "success":
            parsed = safe_parse_json(res["raw_text"])
        
        if parsed:
            fixed_sec = safe_section_fix(parsed, s_title)
            if id_to_url:
                sec_cites = set(re.findall(r'\[(\w+)\]', str(fixed_sec.get("content", ""))))
                # id_to_url ở đây là global_map {id: passage_dict}
                fixed_sec["citations"] = [
                    {"id": cid, "url": id_to_url.get(cid, {}).get("url", "")} 
                    for cid in sec_cites if cid in id_to_url
                ]
            rescue_sections.append(fixed_sec)
        else:
            # Local fallback for this specific section
            rescue_sections.append({
                "title": clean_title_numbering(s_title), 
                "content": f"[Nội dung đang được hệ thống xử lý bổ sung từ nguồn cho mục {s_title}...]", 
                "citations": []
            })

    # Return rescued chapter if we have any valid sections, otherwise use raw fallback
    # V29.2: Exclude placeholder text length from total_len calculation to prevent false success.
    total_len = 0
    for s in rescue_sections:
        c = s.get("content", "")
        if "[Nội dung đang được" not in c:
            total_len += len(c)
            
    if rescue_sections and total_len > 50:
        CONG_VIEC[ctx.ma_cv]["chuong_hoan_thanh"] = CONG_VIEC[ctx.ma_cv].get("chuong_hoan_thanh", 0) + 1
        return {
            "title": clean_title_numbering(chap_title),
            "sections": rescue_sections,
            "status": "rescued"
        }
    else:
        logger.warning(f"{prefix}: [RESCUE FAIL] Gemini output too short or empty ({total_len} chars). Falling back to raw facts.")

    # Fallback Logic (Final attempt - Raw Facts)
    return fallback_raw_facts(chap_info, all_passages)

def parallel_generate(ctx, raw_outline, outline_data):
    """Điều phối biên soạn song song các chương với Context Object (ctx)."""
    cfg = CONFIG_QUY_MO.get(ctx.quy_mo, CONFIG_QUY_MO["tieu_chuan"])
    max_workers = cfg["parallelism"]
    
    final_chapters = []
    # Dùng parallelism động theo quy mô (V17.0+)
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = []
        for idx, chap in enumerate(raw_outline):
            time.sleep(random.uniform(0.1, 0.3)) # Giảm jitter để start nhanh hơn
            futures.append(executor.submit(
                process_chapter_supervisor, ctx, idx, chap, openai_writer, giam_sat_chuong
            ))
        final_chapters = [f.result() for f in futures]
    return final_chapters

@app.route('/admin/update_user', methods=['POST'])
@login_required
def admin_update_user():
    if not current_user.la_admin:
        return jsonify({"error": "Unauthorized"}), 403
    
    data = request.get_json()
    u_id = data.get('id')
    user = NguoiDung.query.get(u_id)
    
    if not user:
        return jsonify({"error": "User not found"}), 404
    
    if 'ho_ten' in data: user.ho_ten = data['ho_ten']
    if 'email' in data: user.email = data['email']
    if 'token' in data: user.token = int(data['token'])
    if 'la_admin' in data: user.la_admin = bool(data['la_admin'])
    
    db.session.commit()
    return jsonify({"success": True})

@app.post("/tao")
def tao_giao_trinh():
    du_lieu = request.get_json(silent=True) or request.form.to_dict()
    tieu_de = (du_lieu.get("tieu_de") or du_lieu.get("title") or "").strip()
    if not tieu_de: return jsonify({"loi": "Thiếu tiêu đề."}), 400
    
    if not is_valid_query(tieu_de) or not is_meaningful(tieu_de):
        return jsonify({
            "status": "INVALID_INPUT",
            "loi": "Nội dung nhập không hợp lệ. Vui lòng nhập từ khóa rõ ràng, có ý nghĩa (chữ, số, không chứa ký tự lạ)."
        }), 400

    ma_cv = str(uuid.uuid4())
    CONG_VIEC[ma_cv] = {"trang_thai": "dang_chay", "tien_do": 0, "tieu_de": tieu_de, "nhat_ky": []}
    
    # Lấy các tham số nâng cao (V31+)
    so_chuong_custom = du_lieu.get("so_chuong_custom")
    danh_sach_chuong = du_lieu.get("danh_sach_chuong")
    che_do = du_lieu.get("che_do", "auto")

    # Xác định phí token (V32+)
    phi_token = 1
    if che_do == "expert": phi_token = 2
    elif che_do == "creative": phi_token = 3

    # Lấy user_id ngay trong request context trước khi chuyển sang background thread
    from flask_login import current_user
    if current_user.is_authenticated:
        # Đặc quyền Admin: Không giới hạn Token (V33+)
        if not current_user.la_admin:
            if current_user.token < phi_token:
                return jsonify({"loi": f"Bạn cần ít nhất {phi_token} tokens cho chế độ này. Vui lòng mua thêm."}), 403
            current_user.token -= phi_token
            db.session.commit()
    u_id = current_user.id if current_user.is_authenticated else None
    
    # Lấy các tham số nâng cao (V31+)
    so_chuong_custom = du_lieu.get("so_chuong_custom")
    danh_sach_chuong = du_lieu.get("danh_sach_chuong")

    def run_pipeline(user_id, so_chuong_custom=None, danh_sach_chuong=None):
        global NEXT_CID
        import time
        start_time = time.time()
        def ghi_nhat_ky(msg):
            ts = datetime.now().strftime("%H:%M:%S")
            elapsed = time.time() - start_time
            log_line = f"[{ts}] {msg} (T+{elapsed:.1f}s)"
            CONG_VIEC[ma_cv]["nhat_ky"].append(log_line)
            logger.info(f"[Job {ma_cv[:8]}] {msg} (T+{elapsed:.1f}s)")
            # Watchdog Check (V22.2)
            if elapsed > 1800: # 30 Minutes Extended Limit
                logger.error(f"WATCHDOG TRIGGERED for {ma_cv}: Elapsed {elapsed:.1f}s > 1800s. Emergency termination/fallback.")
                raise TimeoutError("Pipeline Watchdog limit reached.")

        with app.app_context():
            try:
                from dich_vu.meta_controller import meta_controller_instance
                meta_controller_instance.reset_state()
                
                ghi_nhat_ky(f"Khởi động Pipeline AI cho chủ đề: {tieu_de}")
                from dich_vu.lay_wikipedia import tao_tai_lieu_tu_wikipedia, ekre_discovery_engine
                from dich_vu.lam_sach_van_ban import chia_doan, lam_sach_trang

                quy_mo = du_lieu.get("quy_mo", "tieu_chuan")
                ngon_ngu = du_lieu.get("ngon_ngu", "vi")
                CONG_VIEC[ma_cv]["ngon_ngu"] = ngon_ngu
                logger.info(f"Job {ma_cv}: Pipeline started with SCALE={quy_mo}, LANG={ngon_ngu}, CUSTOM_CH={so_chuong_custom}, MANUAL_LIST={'YES' if danh_sach_chuong else 'NO'} for topic '{tieu_de}'")

                # --- Step 0: Safety Classification (V29 — 3-Layer) ---
                from dich_vu.safety_router import classify_topic, reframe_topic, generate_safe_title, get_block_message
                safety_res = classify_topic(tieu_de, CauHinh.OPENAI_API_KEY)
                safety_class = safety_res.get("classification", "SAFE")
                
                if safety_class == "BLOCK":
                    logger.error(f"Job {ma_cv}: Topic '{tieu_de}' BLOCKED (Layer: {safety_res.get('layer')}). Reason: {safety_res.get('reason')}")
                    block_msg = get_block_message(safety_res)
                    error_text = block_msg["message"] if block_msg else safety_res.get("reason")
                    if block_msg and block_msg.get("suggestion"):
                        error_text += f"\n\n{block_msg['suggestion']}"
                    CONG_VIEC[ma_cv]["loi"] = error_text
                    CONG_VIEC[ma_cv]["trang_thai"] = "that_bai"
                    CONG_VIEC[ma_cv]["loai_loi"] = "safety_block"
                    return
                
                ekre_query = tieu_de
                if safety_class == "REFRAME":
                    ekre_query = reframe_topic(tieu_de)
                    CONG_VIEC[ma_cv]["tieu_de"] = generate_safe_title(tieu_de)
                    ghi_nhat_ky(f"Chủ đề nhạy cảm. Đã chuyển sang phân tích học thuật: {ekre_query}")
                    logger.info(f"[SAFETY] Reframed: '{tieu_de}' → '{ekre_query}'")

                # Bước 1: Khám phá tri thức (EKRE Discovery)
                CONG_VIEC[ma_cv].update({"tien_do": 10, "buoc": "Đang thực hiện EKRE Discovery..."})
                ghi_nhat_ky("Bắt đầu pha Discovery (EKRE Adaptive Harvesting).")
                # Bắt đầu pha Discovery (EKRE Adaptive Harvesting)
                from dich_vu.lay_wikipedia import ekre_discovery_engine
                
                # EKRE trả về passages, candidates, hardened_docs, và xray
                ekre_res = ekre_discovery_engine(
                    ekre_query, 
                    api_keys_list=CauHinh.GEMINI_API_KEYS,
                    quy_mo=quy_mo,
                    api_key_openai=CauHinh.OPENAI_API_KEY,
                    original_topic=tieu_de  # V8.1: Luôn truyền tên gốc để EKRE search Wikipedia chính xác
                )
                passages = ekre_res.get("passages", [])
                candidates = ekre_res.get("candidates", {})
                hardened_docs = ekre_res.get("hardened_docs", [])
                xray = ekre_res.get("xray", {})
                
                xray["safety_class"] = safety_class
                # V29: Lưu danh sách nguồn với đầy đủ metadata + điểm chất lượng
                source_list = []
                for doc in hardened_docs:
                    source_list.append({
                        "title": doc.get("title", "N/A"),
                        "url": doc.get("url", ""),
                        "lang": doc.get("lang", "vi"),
                        "reason": doc.get("subtopic", ""),
                        "quality_score": round(doc.get("quality_score", 0), 2),
                        "relevance_score": round(doc.get("relevance_score", 0), 3),
                        "text_len": len(doc.get("text", "")),
                        "is_low_priority": doc.get("is_low_priority", False),
                    })
                # Sắp xếp theo quality_score giảm dần
                source_list.sort(key=lambda x: x["quality_score"], reverse=True)
                CONG_VIEC[ma_cv]["top_30_links"] = source_list[:30]
                CONG_VIEC[ma_cv]["discovery_xray"] = xray 
                
                # 💎 STRUCTURE LOGGING (V24.1)
                kb_density = len(passages) / 10 # heuristic density per chapter
                logger.info(f"[STRUCTURE] Scale: {quy_mo} | SearchYield: {xray['stats']['filtered']} | Density: {kb_density:.1f}")
                ghi_nhat_ky(f"Discovery hoàn tất. Tìm thấy {len(candidates)} trang nguồn. (X-Ray Yield: {xray['stats']['filtered']}/{xray['stats']['retrieved']})")
                
                # --- PHANH AN TOÀN (RELIABLE SOURCE GATE) ---
                confidence = xray.get("adaptive", {}).get("confidence_score", 0)
                reliable_docs = [
                    d for d in hardened_docs
                    if d.get("quality_score", 0) >= CauHinh.EKRE_MIN_QUALITY_FLOOR
                    and not d.get("is_low_priority", False)
                ]

                reason = None
                if len(reliable_docs) == 0:
                    reason = "NO_RELIABLE_DOCS"
                elif confidence < 0.25:
                    reason = "LOW_CONFIDENCE"

                if reason:
                    logger.warning(f"[NO_RELIABLE] Query: {ekre_query} | Confidence: {confidence} | Reason: {reason}")
                    
                    sorted_docs = sorted(hardened_docs, key=lambda x: x.get("quality_score", 0), reverse=True)
                    preview = []
                    for d in sorted_docs[:5]:
                        preview.append({
                            "title": d.get("title", ""),
                            "similarity": round(d.get("relevance_score", 0), 3),
                            "quality": round(d.get("quality_score", 0), 2),
                            "final_score": round(d.get("quality_score", 0), 3)
                        })
                        
                    CONG_VIEC[ma_cv].update({
                        "trang_thai": "that_bai",
                        "loai_loi": "NO_RELIABLE_SOURCES",
                        "loi": "Không tìm được nguồn tài liệu đủ độ tin cậy để xây dựng nội dung giáo trình.",
                        "chi_tiet": reason,
                        "suggestions": preview,
                        "query_suggestions": [
                            f"{tieu_de} là gì",
                            f"{tieu_de} cơ bản",
                            f"Tổng quan về {tieu_de}"
                        ]
                    })
                    return # Dừng toàn bộ pipeline một cách êm ái
                
                # --- DATA SUFFICIENCY GATE (V31 — Kiểm tra đủ dữ liệu cho quy mô) ---
                from dich_vu.lay_wikipedia import score_knowledge_base
                kb_score = score_knowledge_base(hardened_docs)
                # Ngưỡng tối thiểu theo quy mô (dựa trên target_score trong EKRE)
                SUFFICIENCY_THRESHOLDS = {
                    "can_ban": 15,      # Tối thiểu ~1-2 bài chất lượng
                    "tieu_chuan": 30,   # Tối thiểu ~3 bài chất lượng  
                    "chuyen_sau": 60    # Tối thiểu ~6 bài chất lượng
                }
                SCALE_DOWNGRADE = {
                    "chuyen_sau": "tieu_chuan",
                    "tieu_chuan": "can_ban"
                }
                SCALE_LABELS = {
                    "can_ban": "Căn bản (4-5 chương)",
                    "tieu_chuan": "Tiêu chuẩn (7-10 chương)",
                    "chuyen_sau": "Chuyên sâu (12-20 chương)"
                }
                
                min_threshold = SUFFICIENCY_THRESHOLDS.get(quy_mo, 30)
                
                # --- TOPIC PRESENCE CHECK (V31.8 — Diacritics-aware) ---
                # Loại bỏ từ đệm phổ biến để lấy cụm danh từ cốt lõi
                from dich_vu.lam_sach_van_ban import remove_diacritics
                VIET_STOPWORDS = {
                    "tổng", "quan", "về", "giới", "thiệu", "cơ", "bản", "nâng", "cao",
                    "chuyên", "sâu", "nhập", "môn", "đại", "cương", "khái", "niệm",
                    "của", "và", "các", "trong", "cho", "với", "từ", "đến", "là",
                    "một", "những", "được", "có", "này", "theo", "tại", "trên",
                    # Bản không dấu của stopwords
                    "tong", "ve", "gioi", "co", "ban", "nang",
                    "chuyen", "nhap", "dai", "cuong", "khai",
                    "cua", "va", "cac", "trong", "voi", "tu", "den", "la",
                    "mot", "nhung", "duoc", "nay", "tai"
                }
                topic_lower = tieu_de.lower().strip()
                # Lấy cụm danh từ cốt lõi (bỏ stopwords + từ quá ngắn)
                core_words = [w for w in topic_lower.split() if w not in VIET_STOPWORDS and len(w) > 1]
                core_phrase = " ".join(core_words)
                core_phrase_nodiac = remove_diacritics(core_phrase).lower()
                
                # Sinh tất cả subphrases liên tiếp ≥2 từ (sắp xếp từ dài → ngắn)
                # Ví dụ: "nguyên lí bom nguyên tử" → ["nguyên lí bom nguyên tử", "nguyên lí bom nguyên", "lí bom nguyên tử", "nguyên lí bom", "bom nguyên tử", ...]
                subphrases = []
                if len(core_words) >= 2:
                    for length in range(len(core_words), 1, -1):  # Từ dài đến ngắn
                        for start in range(len(core_words) - length + 1):
                            sp = " ".join(core_words[start:start+length])
                            subphrases.append(sp)
                elif core_phrase:
                    subphrases = [core_phrase]
                
                subphrases_nodiac = [remove_diacritics(sp).lower() for sp in subphrases]
                
                logger.info(f"[DATA-GATE] Core: '{core_phrase}' | Subphrases: {subphrases[:5]}...")
                
                docs_mentioning_topic = 0
                for d in hardened_docs:
                    doc_text = (d.get("text", "") + " " + d.get("title", "")).lower()
                    doc_text_nodiac = remove_diacritics(doc_text)
                    # Match BẤT KỲ subphrase nào (có dấu hoặc không dấu)
                    matched = any(
                        sp in doc_text or sp_nd in doc_text_nodiac
                        for sp, sp_nd in zip(subphrases, subphrases_nodiac)
                    )
                    if matched:
                        docs_mentioning_topic += 1
                
                topic_presence_ratio = docs_mentioning_topic / max(len(hardened_docs), 1)
                avg_sim = xray.get("stats", {}).get("avg_sim", 1.0)
                
                logger.info(
                    f"[DATA-GATE] kb_score={kb_score:.1f} | threshold={min_threshold} | scale={quy_mo} | "
                    f"topic_presence={docs_mentioning_topic}/{len(hardened_docs)} ({topic_presence_ratio:.0%}) | "
                    f"avg_sim={avg_sim:.3f}"
                )
                
                # DỪNG chỉ khi CẢ HAI: (1) phrase không tìm thấy VÀ (2) avg_sim thấp
                # → "phim ma" tìm được "phim kinh dị" (avg_sim cao) → KHÔNG chặn
                # → "phi công" tìm được IT (avg_sim thấp + phrase 0%) → CHẶN
                # V31.8.1: Nới lỏng - Nếu có ít nhất 1 bài khớp topic, chỉ chặn nếu tập tài liệu quá tệ (avg_sim < 0.35)
                is_topic_absent = (docs_mentioning_topic == 0) or (topic_presence_ratio < 0.10 and avg_sim < 0.35)
                if is_topic_absent and avg_sim < 0.45:
                    reason_detail = "TOPIC_NOT_FOUND" if avg_sim < 0.35 else "LOW_RELEVANCE"
                    logger.warning(
                        f"[DATA-GATE] {reason_detail}: '{tieu_de}' (core='{core_phrase}') "
                        f"presence={docs_mentioning_topic}/{len(hardened_docs)}, avg_sim={avg_sim:.3f}"
                    )
                    kb_score = 0  # Ép điểm về 0 để kích hoạt gate bên dưới
                
                if kb_score < min_threshold:
                    # Xác định quy mô khả thi (nếu có)
                    suggested_scale = None
                    if quy_mo in SCALE_DOWNGRADE:
                        lower = SCALE_DOWNGRADE[quy_mo]
                        lower_threshold = SUFFICIENCY_THRESHOLDS[lower]
                        if kb_score >= lower_threshold:
                            suggested_scale = lower
                    
                    # Gợi ý chủ đề thay thế bằng LLM (1 call nhẹ)
                    alt_topics = []
                    try:
                        from openai import OpenAI
                        _client = OpenAI(api_key=CauHinh.OPENAI_API_KEY, max_retries=0)
                        _resp = _client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[{"role": "user", "content": f"""Chủ đề "{tieu_de}" không tồn tại hoặc không có đủ tài liệu học thuật trên Wikipedia.
Hãy gợi ý 4 chủ đề học thuật KHÁC BIỆT HOÀN TOÀN nhưng thuộc cùng lĩnh vực, mà CHẮC CHẮN có bài viết riêng trên Wikipedia tiếng Việt hoặc tiếng Anh.

QUAN TRỌNG:
- KHÔNG gợi ý biến thể của "{tieu_de}" (ví dụ: "Tổng quan về {tieu_de}", "{tieu_de} cơ bản", "Giới thiệu {tieu_de}")
- Chỉ gợi ý các chủ đề CÓ THẬT, phổ biến, có nhiều tài liệu
- Mỗi chủ đề phải là tên riêng của một lĩnh vực/khái niệm cụ thể

Trả về ONLY JSON: {{"topics": ["topic1", "topic2", "topic3", "topic4"]}}"""}],
                            temperature=0.3, timeout=15.0
                        )
                        import json as _json
                        alt_data = _json.loads(_resp.choices[0].message.content)
                        alt_topics = alt_data.get("topics", [])[:4]
                    except Exception as e_alt:
                        logger.warning(f"[DATA-GATE] Alternative topic suggestion failed: {e_alt}")
                        alt_topics = []
                    
                    if topic_presence_ratio < 0.10:
                        error_msg = (
                            f"Chủ đề \"{tieu_de}\" không được tìm thấy trên Wikipedia. "
                            f"Các tài liệu hệ thống thu thập được không hề đề cập đến chủ đề này. "
                            f"Vui lòng kiểm tra lại chính tả hoặc thử một chủ đề khác."
                        )
                    else:
                        error_msg = (
                            f"Chủ đề \"{tieu_de}\" không có đủ dữ liệu chất lượng trên Wikipedia "
                            f"để xây dựng giáo trình ở quy mô {SCALE_LABELS.get(quy_mo, quy_mo)}."
                        )
                    
                    gate_result = {
                        "trang_thai": "that_bai",
                        "loai_loi": "INSUFFICIENT_DATA",
                        "loi": error_msg,
                        "kb_score": round(kb_score, 1),
                        "required_score": min_threshold,
                        "current_scale": quy_mo,
                        "current_scale_label": SCALE_LABELS.get(quy_mo, quy_mo),
                        "query_suggestions": alt_topics,
                    }
                    
                    # Nếu có thể hạ quy mô → gợi ý (KHÔNG tự động hạ, hỏi người dùng)
                    if suggested_scale:
                        gate_result["suggested_scale"] = suggested_scale
                        gate_result["suggested_scale_label"] = SCALE_LABELS.get(suggested_scale, suggested_scale)
                        gate_result["loi"] += (
                            f"\n\nTuy nhiên, dữ liệu hiện có ({kb_score:.0f} điểm) đủ để tạo giáo trình "
                            f"ở mức {SCALE_LABELS.get(suggested_scale, suggested_scale)}."
                        )
                    
                    logger.warning(
                        f"[DATA-GATE] INSUFFICIENT: '{tieu_de}' | score={kb_score:.1f} < {min_threshold} | "
                        f"suggested_scale={suggested_scale} | alt_topics={alt_topics}"
                    )
                    CONG_VIEC[ma_cv].update(gate_result)
                    return  # DỪNG pipeline — chờ người dùng quyết định
                
                # Mọi thứ an toàn, tiếp tục tạo Vector DB
                if not passages:
                    logger.warning(f"Job {ma_cv}: EKRE found no documents. Triggering fallback search...")
                
                # Tạo Vector DB từ EKRE Passages
                passages_db = tao_vector_db(passages, api_key=CauHinh.OPENAI_API_KEY)
                global_map = {p['id']: p for p in passages_db}
                
                # 🛠️ KHỞI TẠO PipelineContext: Trung tâm điều phối Metadata
                ctx = PipelineContext(
                    ma_cv=ma_cv,
                    tieu_de=tieu_de,
                    quy_mo=quy_mo,
                    api_keys_list=CauHinh.GEMINI_API_KEYS,
                    passages_db=passages_db,
                    global_map=global_map,
                    terms=du_lieu.get("terms", []),
                    passages=passages,
                    candidates=candidates,
                    openai_semaphore=OPENAI_SEMAPHORE,
                    safety_class=safety_class,
                    ngon_ngu=ngon_ngu
                )
                prefix = ctx.get_logger_prefix()
                
                # Bước 2: Dàn ý (Outline) & Thuật ngữ
                CONG_VIEC[ctx.ma_cv].update({"tien_do": 30, "buoc": "Đang phân tích tri thức & trích xuất thuật ngữ..."})
                
                try:
                    from dich_vu.openai_da_buoc import (
                        xay_dung_metadata_toan_dien, 
                        trich_xuat_thuat_ngu,
                        nhom_thuat_ngu_va_tao_dan_y,
                        tao_dan_y_tu_passages,
                        xac_dinh_ngan_sach_thuat_ngu,
                        get_structure_config,
                        InsufficientDataError
                    )
                    
                    # V23.5.4: Logging and Budgeting fix (Synchronized with Scales)
                    num_articles = len(ctx.candidates) if ctx.candidates else 1
                    so_chuong_yeu_cau = int(du_lieu.get("so_chuong", 0)) 
                    
                    # 1. Budgeting (Xác định quy mô và định mức thuật ngữ)
                    budget = xac_dinh_ngan_sach_thuat_ngu(num_articles, so_chuong_yeu_cau, quy_mo=ctx.quy_mo)
                    
                    target_ch_log = so_chuong_yeu_cau if so_chuong_yeu_cau > 0 else f"{get_structure_config(ctx.quy_mo)['ch'][0]}-{get_structure_config(ctx.quy_mo)['ch'][1]}"
                    logger.info(f"[STRUCTURE] Scale: {ctx.quy_mo} | TermBudget: {budget['core_count']} | ChapterGoal: {target_ch_log}")
                    
                    ghi_nhat_ky(f"Ngân sách thuật ngữ: {budget['core_count']} core. Mục tiêu: ~{target_ch_log} chương.")

                    # 2. Metadata Builder
                    metadata_list = xay_dung_metadata_toan_dien(ctx.passages)
                    
                    # 2. Term Extraction (Adaptive V23.3) - 🧵 Buffered by Semaphore
                    ghi_nhat_ky(f"Bắt đầu trích xuất thuật ngữ khoa học (Target: {budget['core_count'] + budget['support_count']}).")
                    step_start = time.time()
                    terms_data = trich_xuat_thuat_ngu(
                        ctx.passages, api_key=CauHinh.OPENAI_API_KEY, 
                        target_core=budget["core_count"], 
                        target_support=budget["support_count"],
                        semaphore=ctx.openai_semaphore
                    )
                    logger.info(f"{prefix}: Term Extraction completed in {time.time()-step_start:.2f}s")
                    ctx.terms = terms_data.get("core_terms", []) + terms_data.get("supporting_terms", [])
                    CONG_VIEC[ma_cv]["terms_detail"] = terms_data
                    ghi_nhat_ky(f"Đã trích xuất {len(ctx.terms)} thuật ngữ khoa học.")
                    CONG_VIEC[ctx.ma_cv]["terms_count"] = len(ctx.terms)
                    
                    # 4. Validation Layer (Term Intensity Check)
                    expected_min = budget["core_count"] * 0.6
                    if len(ctx.terms) < expected_min:
                        logger.warning(f"{prefix}: Insufficient core terms ({len(ctx.terms)} < {expected_min}). Triggering fallback.")
                        raise ValueError("Insufficient core terms extracted.")

                    # 5. Advanced Outline Generation (Phases & Uniqueness) - 🧵 Buffered by Semaphore
                    step_start = time.time()
                    outline_data = nhom_thuat_ngu_va_tao_dan_y(
                        terms_data, 
                        api_key=CauHinh.OPENAI_API_KEY, 
                        chu_de=ctx.tieu_de,
                        so_chuong=so_chuong_yeu_cau,
                        quy_mo=ctx.quy_mo,
                        semaphore=ctx.openai_semaphore,
                        ngon_ngu=ngon_ngu,
                        so_chuong_custom=so_chuong_custom,
                        danh_sach_chuong=danh_sach_chuong
                    )
                    actual_chapters = len(outline_data.get("outline", []))
                    logger.info(f"{prefix}: Advanced Outline completed in {time.time()-step_start:.2f}s. Chapters: {actual_chapters}")
                    
                    # 💎 Structural Audit (V24.6)
                    target_ch_audit = so_chuong_yeu_cau if so_chuong_yeu_cau > 0 else (get_structure_config(ctx.quy_mo).get("ch", (4, 8))[0] + get_structure_config(ctx.quy_mo).get("ch", (4, 8))[1]) // 2
                    if actual_chapters < target_ch_audit:
                        logger.warning(f"[STRUCTURE] Job {ma_cv}: Structural Violation (Got {actual_chapters}, Expected {target_ch_audit}).")
                    
                    if not outline_data or not outline_data.get("outline"):
                        raise ValueError("Advanced Outline failed or empty.")
                    
                    # --- PER-CHAPTER OUTLINE GUARD (V23.5.1 Synchronized) ---
                    raw_outline = outline_data.get("outline", [])
                    struct_cfg = get_structure_config(ctx.quy_mo)
                    min_required = struct_cfg["sec"][0] # Lấy số mục tối thiểu từ cấu hình chung
                    
                    thin_chapters = [c.get("title") for c in raw_outline if len(c.get("sections", [])) < min_required]
                    if thin_chapters:
                        logger.warning(f"{prefix}: Found {len(thin_chapters)} thin chapters. Retrying with STRICT_FRAGMENTATION...")
                        retry_outline = nhom_thuat_ngu_va_tao_dan_y(
                            terms_data, api_key=CauHinh.OPENAI_API_KEY, 
                            chu_de=f"{ctx.tieu_de} (STRICT_FRAGMENTATION: Each chapter MUST have {min_required} to {struct_cfg['sec'][1]} sections)",
                            so_chuong=so_chuong_yeu_cau,
                            quy_mo=ctx.quy_mo,
                            semaphore=ctx.openai_semaphore,
                            ngon_ngu=ngon_ngu,
                            so_chuong_custom=so_chuong_custom,
                            danh_sach_chuong=danh_sach_chuong
                        )
                        if retry_outline and retry_outline.get("outline"):
                            outline_data = retry_outline
                            actual_ch_retry = len(outline_data.get("outline", []))
                            logger.info(f"{prefix}: Strict Fragmentation Outline completed in {time.time()-step_start:.2f}s. Chapters: {actual_ch_retry}")

                except Exception as ex_prod:
                    logger.warning(f"{prefix}: Tier 1 Outline failed ({ex_prod}). Attempting Tier 2 (Smart Passages-to-Outline)...")
                    try:
                        # Tier 2: Smart Passages-to-Outline Fallback
                        step_start = time.time()
                        outline_data = tao_dan_y_tu_passages(
                            ctx.tieu_de, ctx.passages, 
                            api_key=CauHinh.OPENAI_API_KEY, 
                            quy_mo=ctx.quy_mo,
                            semaphore=ctx.openai_semaphore
                        )
                        logger.info(f"{prefix}: Tier 2 Smart Outline completed in {time.time()-step_start:.2f}s")
                    except Exception as ex_tier2:
                        logger.warning(f"{prefix}: Tier 2 Outline failed ({ex_tier2}). Falling back to Tier 3 (Legacy Rescue)...")
                        # Tier 3: Legacy Outline Rescue (Oldest Path)
                        top_pass = tim_kiem_vector(ctx.tieu_de, ctx.passages_db, api_key=CauHinh.OPENAI_API_KEY, top_k=60)
                        che_do = du_lieu.get("che_do", "auto")
                        outline_data = openai_tao_dan_y(
                            ctx.tieu_de, top_pass, api_key=CauHinh.OPENAI_API_KEY, 
                            quy_mo=ctx.quy_mo, che_do=che_do, so_chuong_max=8,
                            semaphore=ctx.openai_semaphore
                        )
                
                raw_outline = outline_data.get("outline", [])
                total_sections = sum([len(c.get("sections", [])) for c in raw_outline])
                CONG_VIEC[ma_cv]["tong_chuong"] = len(raw_outline)
                CONG_VIEC[ma_cv]["tong_muc"] = total_sections
                ghi_nhat_ky(f"Dàn ý hoàn tất: {len(raw_outline)} chương, {total_sections} mục con.")
                logger.info(f"{prefix}: Outline created with {len(raw_outline)} chapters and {total_sections} sections.")

                # --- 🚀 LEVEL 3 EXPANSION: Chapter-driven (V17.2 Hardened) ---
                missing_chapter_topics = []
                # Đọc an toàn list hiện tại
                with PASSAGES_LOCK:
                    current_db_snapshot = list(passages_db)
                
                if len(current_db_snapshot) < MAX_TOTAL_PASSAGES:
                    for chap in raw_outline:
                        title = chap.get("title", "")
                        # tim_kiem_vector đã có snapshot nội bộ, không lo race
                        hits = tim_kiem_vector(title, current_db_snapshot, api_key=CauHinh.OPENAI_API_KEY, top_k=2)
                        if not hits or len(hits) < 1:
                            logger.info(f"{prefix}: KB Gap detected for '{title}'. Adding to expansion queue.")
                            missing_chapter_topics.append(title)
                else:
                    logger.info(f"{prefix}: KB Limit reached ({len(current_db_snapshot)}). Skipping expansion.")

                if missing_chapter_topics:
                    CONG_VIEC[ma_cv]["buoc"] = "Đang mở rộng kiến thức theo chương (Expansion)..."
                    from dich_vu.lay_wikipedia import smart_search_crawl
                    # Giới hạn crawl trong 5 docs để an toàn
                    new_docs_raw = smart_search_crawl(missing_chapter_topics[:5])
                    
                    if new_docs_raw:
                        try:
                            # V21.6: Embedded docs with automatic ID assignment
                            logger.info(f"{prefix}: Embedding {len(new_docs_raw)} new docs...")
                            next_id = len(current_db_snapshot) + 1
                            embedded_new = tao_vector_db(new_docs_raw, api_key=CauHinh.OPENAI_API_KEY, start_id=next_id)
                            
                            if embedded_new:
                                # 2. Thread-safe Append
                                with PASSAGES_LOCK:
                                    passages_db.extend(embedded_new)
                                    global_map.update({p['id']: p for p in embedded_new})
                                    ctx.passages_db = passages_db 
                                
                                logger.info(f"{prefix}: Added {len(embedded_new)} embedded docs to KB.")
                        except Exception as ex_embed:
                            logger.error(f"{prefix}: Expansion embedding failed: {ex_embed}")

                # Bước 3: Biên soạn nội dung (V17.0+ Turbo Resilience)
                CONG_VIEC[ma_cv].update({"tien_do": 50, "buoc": "Đang biên soạn nội dung (Parallel Writing)..."})
                ghi_nhat_ky("Bắt đầu giai đoạn biên soạn nội dung song song (Multi-threaded Micro-Writer).")
                final_chapters = parallel_generate(ctx, raw_outline, outline_data)
                ghi_nhat_ky(f"Biên soạn xong {len(final_chapters)} chương.")

                # --- V33: CHAPTER SUMMARIES & GLOSSARY ---
                CONG_VIEC[ma_cv].update({"tien_do": 70, "buoc": "Đang sinh tóm tắt chương & bảng thuật ngữ..."})
                from dich_vu.openai_da_buoc import sinh_tom_tat_chuong, sinh_bang_thuat_ngu
                
                # Sinh tóm tắt cho từng chương (song song)
                from concurrent.futures import ThreadPoolExecutor as _TPE
                def _gen_summary(chap):
                    sections_text = "\n".join(sec.get("content", "") for sec in chap.get("sections", []))
                    return sinh_tom_tat_chuong(tieu_de, chap.get("title", ""), sections_text, CauHinh.OPENAI_API_KEY, OPENAI_SEMAPHORE)
                
                with _TPE(max_workers=4) as _ex:
                    summaries = list(_ex.map(_gen_summary, final_chapters))
                for i, chap in enumerate(final_chapters):
                    chap["summary"] = summaries[i] if i < len(summaries) else ""
                ghi_nhat_ky(f"Đã sinh tóm tắt cho {sum(1 for s in summaries if s)} chương.")
                
                # Sinh bảng thuật ngữ (glossary)
                glossary = sinh_bang_thuat_ngu(ctx.terms, tieu_de, CauHinh.OPENAI_API_KEY, OPENAI_SEMAPHORE)
                ghi_nhat_ky(f"Bảng thuật ngữ: {len(glossary)} định nghĩa.")
                CONG_VIEC[ma_cv]["glossary"] = glossary

                ghi_nhat_ky("Bắt đầu hậu xử lý trích dẫn.")

                # Bước 4: Hậu xử lý trích dẫn & Tham khảo (V18.6 Sequential)
                all_original_passages = {str(p['id']): p for p in passages_db}
                url_to_new_id = {}
                ordered_refs = []
                
                # Scan 1: Xây dựng bản đồ trích dẫn duy nhất theo thứ tự xuất hiện
                for chap in final_chapters:
                    for sec in chap.get("sections", []):
                        found_ids = re.findall(r'\[(\w+)\]', sec.get("content", ""))
                        for oid in found_ids:
                            if oid in all_original_passages:
                                p = all_original_passages[oid]
                                url = p.get('url', '')
                                if url and url not in url_to_new_id:
                                    new_id = len(url_to_new_id) + 1  # V31.6: Số liên tiếp 1,2,3... kiểu int
                                    url_to_new_id[url] = new_id
                                    ordered_refs.append({"id": new_id, "url": url, "title": p.get('title', 'Nguồn')})

                # Scan 2: Cập nhật mã trích dẫn mới vào nội dung và metadata
                for chap in final_chapters:
                    for sec in chap.get("sections", []):
                        content = sec.get("content", "")
                        found_ids = re.findall(r'\[(\w+)\]', content)
                        sec_citations = []
                        added_urls = set()
                        
                        # Cập nhật mã trích dẫn trong văn bản
                        for oid in set(found_ids):
                            if oid in all_original_passages:
                                url = all_original_passages[oid].get('url', '')
                                if url in url_to_new_id:
                                    nid = url_to_new_id[url]
                                    title = all_original_passages[oid].get('title', 'Nguồn').replace('"', '&quot;')
                                    html_tag = f'<sup class="citation"><a href="{url}" title="{title}" target="_blank" rel="noopener noreferrer">[{nid}]</a></sup>'
                                    content = content.replace(f"[{oid}]", html_tag)
                                    if url not in added_urls:
                                        node = next(r for r in ordered_refs if r["url"] == url)
                                        sec_citations.append(node)
                                        added_urls.add(url)
                            else:
                                # --- FINAL SANITIZATION ---
                                # Gỡ bỏ các token [ID] tự bịa / lỗi không có trong hệ thống 
                                content = re.sub(rf"\[{oid}\]", "", content)
                        
                        sec["content"] = content
                        sec["citations"] = sec_citations

                # --- V33: GROUNDING SCORE CALCULATION ---
                grounding_stats = {"chapters": [], "overall": {}}
                total_paras_all = 0
                grounded_paras_all = 0

                for chap in final_chapters:
                    chap_title_gs = chap.get("title", "")
                    chap_total = 0
                    chap_grounded = 0
                    for sec in chap.get("sections", []):
                        paragraphs = [p.strip() for p in sec.get("content", "").split("\n") if p.strip()]
                        for para in paragraphs:
                            chap_total += 1
                            if re.search(r'<sup class="citation">', para) or re.search(r'\[\d+\]', para):
                                chap_grounded += 1
                    ratio_gs = (chap_grounded / chap_total * 100) if chap_total > 0 else 0
                    grounding_stats["chapters"].append({
                        "title": chap_title_gs, "total": chap_total,
                        "grounded": chap_grounded, "ratio": round(ratio_gs, 1)
                    })
                    total_paras_all += chap_total
                    grounded_paras_all += chap_grounded

                overall_ratio = (grounded_paras_all / total_paras_all * 100) if total_paras_all > 0 else 0
                grounding_stats["overall"] = {
                    "total": total_paras_all, "grounded": grounded_paras_all,
                    "ratio": round(overall_ratio, 1)
                }
                CONG_VIEC[ma_cv]["grounding"] = grounding_stats
                ghi_nhat_ky(f"Grounding Score: {overall_ratio:.1f}% ({grounded_paras_all}/{total_paras_all} đoạn có nguồn)")


                # 3. Đóng gói Giáo trình Final
                from dich_vu.kiem_tra_cau_truc_json import clean_title_numbering
                final_book = {"title": tieu_de, "sections": [], "references": []}
                book_export = {"title": tieu_de, "chapters": []}

                for chap in final_chapters:
                    c_title = clean_title_numbering(chap.get("title", "Không tên"))
                    new_chap = {"title": c_title, "summary": chap.get("summary", ""), "sections": []}
                    final_book["sections"].append({"title": c_title, "is_chapter": True, "content": ""})
                    
                    for sec in chap.get("sections", []):
                        s_title = clean_title_numbering(sec.get("title", "Mục"))
                        s_content = sec.get("content", "")
                        new_chap["sections"].append({"title": s_title, "content": s_content, "citations": sec.get("citations", [])})
                        final_book["sections"].append({"title": s_title, "is_chapter": False, "content": s_content})
                    book_export["chapters"].append(new_chap)

                # Bước 5: Xuất bản PDF/DOCX
                CONG_VIEC[ma_cv].update({"tien_do": 85, "buoc": "Đang đóng gói tài liệu..."})
                all_refs = ordered_refs
                ket_qua = {"topic": tieu_de, "book_vi": book_export, "references": all_refs, "ui_book": final_book, "glossary": glossary, "grounding": grounding_stats}
                
                # --- PHIÊN BẢN SẠCH (KHÔNG TRÍCH DẪN) - V23.5 ---
                import copy
                def strip_citations(text):
                    if not text: return ""
                    # Loại bỏ thẻ <sup> trích dẫn HTML
                    clean = re.sub(r'<sup class="citation">.*?</sup>', '', text)
                    # Loại bỏ các tag [ID] thô nếu còn sót
                    clean = re.sub(r'\[\w+\]', '', clean)
                    return clean

                book_plain = copy.deepcopy(book_export)
                for chap in book_plain.get("chapters", []):
                    for sec in chap.get("sections", []):
                        sec["content"] = strip_citations(sec.get("content", ""))
                        sec["citations"] = [] # Xóa metadata trích dẫn

                ket_qua_plain = {
                    "topic": tieu_de, 
                    "book_vi": book_plain, 
                    "references": [], # Xóa danh mục tham khảo
                    "terms": ket_qua.get("terms", []) # Giữ lại thuật ngữ
                }

                p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}.json")
                p_docx = os.path.join(CauHinh.THU_MUC_DOCX, f"{ma_cv}.docx")
                p_pdf = os.path.join(CauHinh.THU_MUC_PDF, f"{ma_cv}.pdf")
                
                p_docx_plain = os.path.join(CauHinh.THU_MUC_DOCX, f"{ma_cv}_plain.docx")
                p_pdf_plain = os.path.join(CauHinh.THU_MUC_PDF, f"{ma_cv}_plain.pdf")

                # Xuất các phiên bản
                _luu_json(ket_qua, p_json)
                xuat_docx(ket_qua, p_docx); xuat_pdf(ket_qua, p_pdf)
                xuat_docx(ket_qua_plain, p_docx_plain); xuat_pdf(ket_qua_plain, p_pdf_plain)

                CONG_VIEC[ma_cv].update({
                    "trang_thai": "hoan_thanh", "tien_do": 100, "nguon": all_refs,
                    "tai_docx": f"/tai/docx/{ma_cv}", 
                    "tai_pdf": f"/tai/pdf/{ma_cv}",
                    "tai_docx_plain": f"/tai/docx/{ma_cv}_plain",
                    "tai_pdf_plain": f"/tai/pdf/{ma_cv}_plain",
                    "giam_sat": {"chapters": len(final_chapters), "circuit_breaker": getattr(ctx, 'use_gemini_only', False)}
                })
                # Bước 6: Kiểm định cuối cùng (Final Audit for Scale Compliance)
                from dich_vu.gemini_giam_sat import giam_sat_quy_mo
                audit_quy_mo = giam_sat_quy_mo(tieu_de, final_chapters, quy_mo, api_keys=CauHinh.GEMINI_API_KEYS)
                if audit_quy_mo.get("status") == "fail":
                    # V23.5.3 Hardening: Chuyển từ Crash sang Warning để bảo vệ Pipeline
                    logger.warning(f"Quy mô chưa đạt kỳ vọng: {audit_quy_mo.get('issues')}")
                    ghi_nhat_ky(f"Cảnh báo quy mô: {audit_quy_mo.get('status')}. Vẫn tiếp tục đóng gói bản tốt nhất.")
                
                ghi_nhat_ky(f"Quy trình hoàn tất thành công trong {time.time() - start_time:.1f} giây.")
                logger.info(f"Job {ma_cv}: COMPLETED in {time.time() - start_time:.2f} seconds.")

                # Lưu lịch sử database (Hotfix V18.9: Bỏ qua nếu là khách lẻ và DB chưa migrate)
                try:
                    from mo_hinh import LichSuGiaoTrinh
                    with app.test_request_context('/'):
                        noi_dung_html = render_template("result.html", ma_cv=ma_cv, thong_tin=CONG_VIEC[ma_cv], book=book_export, references=all_refs)
                    
                    if user_id:
                        # Tính tổng số ký tự nội dung giáo trình (V32 Fix)
                        tong_ky_tu = 0
                        for chap in book_export.get("chapters", []):
                            for sec in chap.get("sections", []):
                                tong_ky_tu += len(sec.get("content", ""))
                        
                        history_entry = LichSuGiaoTrinh(
                            nguoi_dung_id=user_id, 
                            chu_de=tieu_de, 
                            noi_dung_html=noi_dung_html, 
                            duong_dan_file=p_pdf, 
                            da_xuat_file=True,
                            do_dai_ky_tu=tong_ky_tu
                        )
                        db.session.add(history_entry)
                        db.session.commit()
                except Exception as db_err:
                    db.session.rollback()
                    logger.error(f"DB History Error: {db_err}")

            except Exception as e:
                logger.error(f"Job {ma_cv} failed: {e}\n{traceback.format_exc()}")
                CONG_VIEC[ma_cv].update({"trang_thai": "that_bai", "loi": str(e)})
            finally:
                # V21.6 Hardening: Close DB session and ensure status isn't stuck
                db.session.remove()
                if ma_cv in CONG_VIEC and CONG_VIEC[ma_cv]["trang_thai"] == "dang_chay":
                    CONG_VIEC[ma_cv]["trang_thai"] = "that_bai"
                    CONG_VIEC[ma_cv]["loi"] = "Pipeline terminated unexpectedly."

    # Chạy Background
    import threading
    threading.Thread(target=run_pipeline, args=(u_id, so_chuong_custom, danh_sach_chuong)).start()
    return jsonify({"ma_cv": ma_cv, "trang_thai": "dang_chay"})

@app.get("/trang_thai/<ma_cv>")
def trang_thai(ma_cv):
    data = CONG_VIEC.get(ma_cv, {"trang_thai": "khong_tim_thay"})
    # V30 Fix: Convert numpy types to native Python types for JSON serialization
    return jsonify(json.loads(json.dumps(data, default=_json_safe_default)))

def sanitize_filename(filename):
    """Làm sạch tên file, giữ lại tiếng Việt có dấu (modern browsers support it)."""
    return re.sub(r'[\\/*?:"<>|]', '', filename).strip()

@app.get("/tai/<loai>/<ma>")
def tai_file(loai, ma):
    # Hỗ trợ phiên bản 'plain' bằng cách bóc tách hậu tố để tra cứu info (V23.5.2)
    ma_goc = ma.replace("_plain", "")
    info = CONG_VIEC.get(ma_goc)
    if not info: return "Not found", 404
    
    folder = CauHinh.THU_MUC_PDF if loai == "pdf" else CauHinh.THU_MUC_DOCX
    ext = "pdf" if loai == "pdf" else "docx"
    path = os.path.join(folder, f"{ma}.{ext}")
    
    if os.path.exists(path):
        # Đặt tên file theo chủ đề
        tieu_de = info.get("tieu_de", "giao_trinh")
        filename = f"{sanitize_filename(tieu_de)}.{ext}"
        return send_file(path, as_attachment=True, download_name=filename)
    
    return "File not found on server", 404

@app.get("/tai/zip/<ma>")
def tai_zip(ma):
    """Tải cả Word + PDF trong 1 file ZIP (V32 - Bundle Export)."""
    is_plain = "_plain" in ma
    ma_goc = ma.replace("_plain", "")
    info = CONG_VIEC.get(ma_goc)
    if not info:
        return "Not found", 404

    tieu_de = sanitize_filename(info.get("tieu_de", "giao_trinh"))
    suffix = "_plain" if is_plain else ""
    
    # Xác định đường dẫn file
    p_docx = os.path.join(CauHinh.THU_MUC_DOCX, f"{ma_goc}{suffix}.docx")
    p_pdf = os.path.join(CauHinh.THU_MUC_PDF, f"{ma_goc}{suffix}.pdf")
    
    if not os.path.exists(p_docx) and not os.path.exists(p_pdf):
        return "Files not found on server", 404
    
    # Tạo ZIP trong bộ nhớ (RAM) để tránh ghi đĩa
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        if os.path.exists(p_docx):
            zf.write(p_docx, f"{tieu_de}.docx")
        if os.path.exists(p_pdf):
            zf.write(p_pdf, f"{tieu_de}.pdf")
    
    zip_buffer.seek(0)
    
    label = "ban_sach" if is_plain else "trich_dan"
    zip_filename = f"{tieu_de}_{label}.zip"
    
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name=zip_filename
    )

@app.get("/ket_qua/<ma_cv>")
def ket_qua(ma_cv):
    thong_tin = CONG_VIEC.get(ma_cv)
    if not thong_tin: return "Not found", 404
    if thong_tin.get("trang_thai") == "hoan_thanh":
        p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}.json")
        if os.path.exists(p_json):
            with open(p_json, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return render_template("result.html", ma_cv=ma_cv, thong_tin=thong_tin, book=data.get('book_vi', {}), references=data.get('references', []), glossary=data.get('glossary', []), grounding=data.get('grounding', {}))
    return render_template("result.html", ma_cv=ma_cv, thong_tin=thong_tin)

# --- V33: EXPORT RIÊNG GLOSSARY & SUMMARY ---
@app.get("/tai/glossary/<ma>")
def tai_glossary(ma):
    """Xuất file DOCX chỉ chứa Bảng thuật ngữ."""
    p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma}.json")
    if not os.path.exists(p_json): return "JSON not found", 404

    with open(p_json, 'r', encoding='utf-8') as f:
        data = json.load(f)

    glossary = data.get("glossary", [])
    if not glossary: return "No glossary data", 404

    # Lấy tên chủ đề từ CONG_VIEC hoặc từ JSON
    info = CONG_VIEC.get(ma, {})
    tieu_de = info.get("tieu_de") or data.get("topic", "Giáo trình")

    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    for _ in range(3): doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("BẢNG THUẬT NGỮ")
    run_t.font.size = Pt(20)
    run_t.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sub.add_run(tieu_de)
    run_s.font.size = Pt(14)
    run_s.italic = True

    doc.add_page_break()

    # Glossary entries
    for item in glossary:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run_term = p.add_run(f"{item.get('term', '')}: ")
        run_term.font.size = Pt(13)
        run_term.bold = True
        run_def = p.add_run(item.get('definition', ''))
        run_def.font.size = Pt(13)

    # Save to BytesIO
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True,
                     download_name=f"{sanitize_filename(tieu_de)} - Bảng thuật ngữ.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.get("/tai/summary/<ma>")
def tai_summary(ma):
    """Xuất file DOCX chỉ chứa Tóm tắt các chương."""
    p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma}.json")
    if not os.path.exists(p_json): return "JSON not found", 404

    with open(p_json, 'r', encoding='utf-8') as f:
        data = json.load(f)

    book = data.get("book_vi", {})
    chapters = book.get("chapters", [])
    has_any = any(ch.get("summary") for ch in chapters)
    if not has_any: return "No summary data", 404

    # Lấy tên chủ đề từ CONG_VIEC hoặc từ JSON
    info = CONG_VIEC.get(ma, {})
    tieu_de = info.get("tieu_de") or data.get("topic", "Giáo trình")

    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    for _ in range(3): doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("TÓM TẮT CÁC CHƯƠNG")
    run_t.font.size = Pt(20)
    run_t.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sub.add_run(tieu_de)
    run_s.font.size = Pt(14)
    run_s.italic = True

    doc.add_page_break()

    # Chapter summaries
    for idx, ch in enumerate(chapters, 1):
        h = doc.add_heading(f"Chương {idx}: {ch.get('title', '')}", level=2)
        summary = ch.get("summary", "")
        if summary:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(summary)
            run.font.size = Pt(13)
        else:
            p = doc.add_paragraph("Chưa có tóm tắt cho chương này.")
            p.runs[0].italic = True

    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True,
                     download_name=f"{sanitize_filename(tieu_de)} - Tóm tắt chương.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")



def cleanup_old_jobs():
    """
    Tự động dọn dẹp các Job cũ để tránh chiếm dụng RAM vô thời hạn (V18.3).
    Chạy định kỳ mỗi 15 phút, xóa job đã xong/thất bại nếu danh sách quá dài (>50).
    """
    while True:
        try:
            time.sleep(900) # 15 minutes
            to_delete = []
            for ma_cv, info in list(CONG_VIEC.items()):
                if info.get("trang_thai") in ["hoan_thanh", "that_bai"]:
                    to_delete.append(ma_cv)
            
            # Chỉ cleanup nếu tổng số job vượt ngưỡng an toàn
            if len(CONG_VIEC) > 50:
                for ma_cv in to_delete[:20]:
                    CONG_VIEC.pop(ma_cv, None)
                if to_delete:
                    logger.info(f"[Cleanup] Pruned old jobs. Current count: {len(CONG_VIEC)}")
        except Exception as e:
            logger.error(f"[Cleanup Error] {e}")

# Khởi chạy thread dọn dẹp ngầm
threading.Thread(target=cleanup_old_jobs, daemon=True).start()

if __name__ == "__main__":
    with app.app_context(): db.create_all()
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", 5000)), debug=True)
